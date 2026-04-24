"""Generates the Word document with 54 jury Q&A + practical advice for ONCS 2026."""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Antigravity\__SKILLS\IoanSiAlbert\docs\Intrebari_Juriu_ONCS_2026.docx"

NAVY = RGBColor(0x0F, 0x13, 0x1F)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
CYAN = RGBColor(0x00, 0x88, 0xAA)

# ----------------------------- DATA -----------------------------

SET1 = [
    ("Î1. De ce ați ales distanța de 7 milioane km pentru Dyson Ring?",
     "La 7M km de Soare, intensitatea solară este ~6,22 × 10⁵ W/m² — de 457 de ori mai mare decât la Pământ (1361 W/m²). Parker Solar Probe (NASA, 2018) a ajuns deja la 6,1M km, demonstrând că e fezabil tehnologic. Mai aproape = flux termic inabsorbabil; mai departe = pierdem avantajul concentrației."),
    ("Î2. Cum ați calculat puterea livrată de 1,6 TW?",
     "(1) 30 oglinzi × π·(500 m)² = 2,36 × 10⁷ m² arie totală.  (2) × 6,22 × 10⁵ W/m² = 1,46 × 10¹³ W incident.  (3) × 90% reflectivitate = 1,32 × 10¹³ W.  (4) ÷ 4 colectori × 43% (lumină→laser) = 1,42 × 10¹² W laser/satelit.  (5) 4 sateliți grupați 2+2 pe L4/L5 → 2,83 × 10¹² W / nod.  (6) × 60% (laser→electricitate) × 47% utilizat per centrală = 800 GW/pol × 2 = 1,6 TW."),
    ("Î3. De ce YBCO și nu NbTi (ca la CERN)?",
     "YBCO are T_critică = 93 K → răcire cu azot lichid (77 K). NbTi are T_critică = 9 K, necesită heliu lichid (4 K) — criogenie de ~10× mai grea și mai scumpă. Pentru o structură suspendată de baloane la 30 km, fiecare kilogram contează. YBCO = sistemul criogenic realist."),
    ("Î4. Ce avantaj au punctele L4/L5 față de orbita geostaționară (GEO)?",
     "L4 și L5 sunt puncte de echilibru gravitațional stabil Soare-Pământ. O stație acolo nu consumă combustibil pentru menținerea poziției. GEO necesită corecții orbitale periodice. În plus, L4/L5 oferă linie de vedere continuă spre cei doi poli."),
    ("Î5. Cum stau baloanele cu 11 tone la 30 km altitudine?",
     "Principiul lui Arhimede: F_A = ρ_aer × V × g. La 30 km densitatea aerului este ~0,018 kg/m³ (față de 1,225 la sol). Pentru a ridica 11 tone avem nevoie de ~610.000 m³ de heliu — mari, dar fezabile. Project Loon (Google) opera la altitudini similare cu baloane de ~15 tone."),
    ("Î6. De ce verde pentru laser Soare→L4/L5 și albastru pentru L4/L5→Pământ?",
     "Două raționamente: (1) lungimi de undă optimizate — verdele (532 nm) are atenuare minimă în spațiu, albastrul (450 nm) pătrunde atmosferei mai bine la coborâre spre poli; (2) cod de culoare diferențiat pentru monitorizare vizuală — fiecare segment al rețelei e identificabil imediat."),
    ("Î28. De ce oglinzi reflective și nu panouri solare direct la 7M km?",
     "Panourile solare la 7M km primesc flux de 622 kW/m² — asta arde orice celulă PV existentă. Oglinzile reflectă 90% din lumină fără absorbție termică semnificativă, rămân reci. Conversia lumină→electricitate se face mai departe (la colectori, la 10M km unde fluxul e mai mic) cu radiatoare termice dedicate. Separarea termică = supraviețuirea sistemului."),
    ("Î29. Cum țineți cele 30 de oglinzi pe orbită — fără propulsie constantă?",
     "Orbită Kepleriană naturală la 7M km cu perioada ~11 zile. Oglinzile stau pe această orbită pasiv. Corecții minore de atitudine prin rotori de moment (reaction wheels) + thrusters ionici pentru station-keeping (~10 m/s delta-v pe an — minim)."),
    ("Î30. Ce se întâmplă cu laserul în timpul nopții polare (6 luni de întuneric)?",
     "Centralele polare funcționează EXACT atunci — asta e genial. Polul Nord are întuneric solar iarna, dar laserul vine de la L4/L5 (situate în orbita Pământului la 60°), mereu disponibil indiferent de anotimp. Polii au priză permanentă, neafectați de unghiul Pământ-Soare."),
    ("Î31. Cum calibrați iris-ul captor L4/L5 să prindă laserul la 150M km distanță?",
     "Divergența laser verde (532 nm) la 150M km ≈ spot de ~1 km. Iris-ul are Ø ~100 m și petale active care se reorientează pe baza: (a) unui semnal pilot laser co-axial (urmărire activă); (b) feedback-ului de intensitate pe hub-ul central — dacă intensitatea scade, petalele corectează. Timp de răspuns <1 secundă per corecție."),
    ("Î32. Ce materiale folosiți pentru oglinzi? Nu se deformează termic?",
     "Substrat aluminiu lustruit depus pe beriliu (la fel ca James Webb). Beriliu are coeficient de expansiune termică foarte scăzut + densitate mică (1,85 g/cm³). La 90% reflectivitate, absorbția de 10% e eliminată prin radiatoare pasive pe reversul oglinzii. Deformare termică calculată: <λ/10 la lungimea de undă vizibilă — acceptabil optic."),
    ("Î33. Care e densitatea de curent prin cablul YBCO la 800 GW?",
     "La tensiune HVDC industrială de 1 MV: I = P/U = 800 × 10⁹ / 10⁶ = 800 kA. YBCO industrial suportă j_c ~ 10⁶ A/cm². Necesar teoretic: arie transversală ~0,8 cm². Cu factor de siguranță 10× și redundanță → cablu compus din multe fibre YBCO într-o matrice stabilizatoare (cupru + stainless steel)."),
]

SET2 = [
    ("Î7. Ce problemă rezolvă P.O.L.A.R.I.S.?",
     "Umanitatea consumă ~29.470 TWh electricitate/an (≈ 3,36 TW continuu) din care 60% vine din combustibili fosili. Panourile solare la sol au factor de capacitate ~20% (noaptea, vremea, latitudine). P.O.L.A.R.I.S. livrează 1,6 TW continuu, 24/7, curat — adică 48% din electricitatea globală de la o singură sursă."),
    ("Î8. De ce este ideea VOASTRĂ originală? Există deja concepte similare.",
     "SBSP (Space-Based Solar Power) există ca idee din anii '60 (Glaser, NASA). Dar P.O.L.A.R.I.S. adaugă 3 inovații arhitecturale: (1) separarea colectorilor de emițători — Dyson Ring reflectă, sateliții convertesc, reduce complexitatea termică; (2) relay L4/L5 — nu transmitem direct Soare→Pământ (impracticabil la 150M km cu pierderi pătratice), ci prin noduri Lagrange stabile; (3) cabluri YBCO între stații stratosferice și sol — bypass la problema transmisiei atmosferice wireless."),
    ("Î9. Ce e Scala Kardashev și unde se încadrează proiectul vostru?",
     "Kardashev (1964) clasifică civilizațiile după consum energetic. Tip I = stăpânește toată energia planetei (~10¹⁶ W). Tip II = toată energia stelei (~10²⁶ W). Tip III = galactică. Suntem la Kardashev ~0,73 azi. P.O.L.A.R.I.S. ne urcă la ~0,76. Este primul pas arhitectural spre Tip I."),
    ("Î10. Ce se întâmplă cu surplusul de energie?",
     "Din cei 2,83 × 10¹² W disponibili per nod L4/L5, folosim ~47% pentru cei 800 GW. Restul 53% (~1,5 × 10¹² W) alimentează extensia viitoare Star Power Grid — transmisie prin microunde din orbita GEO spre zone izolate fără infrastructură terestră (Africa sub-sahariană, insulele Pacificului, Arctica)."),
    ("Î11. Care e următorul pas după ONCS?",
     "(a) Simulare numerică mai detaliată a iris-ului captor din L4/L5 (petale trapezoidale); (b) calculul riguros al secțiunii transversale a cablului YBCO pentru 800 GW; (c) analiza stabilității orbitale a celor 30 oglinzi în sector eliptic. Obiectiv concret: extinderea fișei proiectului până în septembrie 2026 pentru posibilă prezentare EUCYS."),
    ("Î12. Pentru cine e util proiectul? Nu va dura 100 de ani?",
     "Util acum — ca framework conceptual pentru tranziția energetică. Dyson Swarms și SBSP sunt discutate în știință, dar rar se vede o arhitectură completă de la Soare la priză. Proiectul nostru închide lanțul. Realizarea fizică? 30–50 de ani pentru un prototip parțial. Dar ideile trebuie testate conceptual acum."),
    ("Î34. Ce impact de mediu are construcția P.O.L.A.R.I.S.?",
     "Construcția implică mii de lansări spațiale, care produc CO₂. Calculul: emisii construcție (~10⁸ tone CO₂ estimat pentru 30 ani) vs emisii evitate (fiecare an funcționare = 1,6 TW × 8760 h × 0,5 kg CO₂/kWh pentru fosili = ~7 × 10⁹ tone CO₂/an evitate). Break-even: ~5 ani de operare. Pe 100 de ani, beneficiu net ~700 × 10⁹ tone CO₂ evitate."),
    ("Î35. Cine ar trebui să controleze P.O.L.A.R.I.S. în operație?",
     "Modelul ITER (reactorul de fuziune internațional): consorțiu multinațional al ONU cu participare obligatorie a puterilor tehnologice majore. Energia = bun comun global. O țară proprietar unic = risc geopolitic maxim. Operațional: centru neutru (Elveția), stații de comandă redundante pe 4 continente."),
    ("Î36. De ce poli și nu ecuator pentru centralele terestre?",
     "Două motive: (1) unghiul laserului — L4/L5 sunt în orbita eliptică a Pământului, iar polii sunt zone slab electrificate, deci oportunitate mai mare de impact; (2) atmosferă mai rară la latitudini înalte + aer mai rece — pierderi mai mici la trecerea laserului spre stațiile stratosferice."),
    ("Î37. Cât ar costa proiectul comparativ cu bugetul global?",
     "Estimare: ~5–10 trilioane USD pentru construcția completă (30 de ani). Investiția globală în energie 2023: ~2,8 trilioane USD/an. P.O.L.A.R.I.S. ≈ 3 ani de investiție globală actuală, distribuit pe 30 ani = ~10% din bugetul energetic anual. Comparativ: Programul Apollo = ~0,4% din PIB-ul SUA/an. P.O.L.A.R.I.S. ar fi similar ca fracțiune globală."),
    ("Î38. Cum se încadrează în angajamentele climatice (Paris, Net Zero 2050)?",
     "Acordul Paris limitează încălzirea <1,5°C — necesită eliminarea ~80% din combustibili fosili până în 2050. P.O.L.A.R.I.S. singur acoperă 48% din electricitate, adică ~60% din decarbonizarea sectorului energetic. Net Zero 2050 devine realist DOAR cu astfel de proiecte megascale, nu doar cu panouri solare pe acoperiș."),
    ("Î39. Care e relația cu ITER și fuziunea nucleară?",
     "Complementare, nu concurente. ITER (fuziune): ~2050–2080 pentru primul reactor comercial. P.O.L.A.R.I.S.: ~2060–2100 pentru primul segment operațional. Ambele trebuie dezvoltate în paralel — fuziunea asigură baza la sol (uzine, orașe); P.O.L.A.R.I.S. asigură capacitate globală + zone izolate. Fuziunea nu rezolvă problema distribuției spre zone slab conectate."),
]

SET3 = [
    ("Î13. Ioan, ce a făcut Albert în proiect?",
     "(Ioan răspunde) Albert s-a ocupat de partea vizuală și narativă — posterul 80×120 cm, scriptul pentru video-ul Kling de 16 s, design-ul paginii polaris-bears.ro. Tot Albert a fundamentat comparația cu Scala Kardashev și a scris secțiunile de viziune din fișă. Fără el, proiectul ar fi fost matematic corect dar ilizibil."),
    ("Î14. Albert, ce a făcut Ioan în proiect?",
     "(Albert răspunde) Ioan s-a ocupat de calculul fizic — a derivat lanțul de eficiențe (90% → 43% → 60%) până la 1,6 TW. A ales parametrii tehnici (distanța 7M km, 30 oglinzi, YBCO vs NbTi) și a verificat fiecare cifră cu legea inversului pătratului. Fără el, numerele nu s-ar închide."),
    ("Î15. Care a fost cel mai mare dezacord tehnic între voi?",
     "(Amândoi) Dezbatere serioasă despre numărul de oglinzi. Ioan voia 20 pentru simetrie matematică, Albert voia 60 pentru redundanță. Compromisul: 30 oglinzi grupate în 4 sectoare, fiecare sector asignat unui colector — satisface și rigoarea matematică, și redundanța."),
    ("Î16. Dacă unul din voi nu ar fi fost în echipă, mai funcționa proiectul?",
     "(Amândoi) Nu. Ioan fără Albert = caiet plin de formule pe care nu le poate vedea un juriu. Albert fără Ioan = poster frumos fără substanță științifică. Complementaritatea e structurală, nu întâmplătoare."),
    ("Î17. Cum comunicați când lucrați la proiect?",
     "Întâlniri fizice săptămânale + mesagerie zilnică. Toate deciziile tehnice mari le luăm împreună, pe tablă, cu calcule scrise. Nu acceptăm „ok, fă tu tu\" pentru deciziile de arhitectură — doar pentru execuție."),
    ("Î18. Cine a avut ideea inițială?",
     "A pornit dintr-o discuție despre cât de irațional e că Pământul primește doar o miliardime din energia Soarelui. Ne-am întrebat: ce-ar fi să captăm fracțiuni mult mai mari? Ideea a evoluat organic — nu e „a lui\" nimănui, e a echipei. (De acolo și Bears — împreună, niciodată singur.)"),
    ("Î19. Care a fost cea mai mare provocare?",
     "(Ioan) Să demonstrez că distanța de 7M km e sustenabilă termic — a trebuit să revizitez datele Parker Solar Probe.\n(Albert) Să fac iris-ul captor L4/L5 vizual corect — am redesenat petalele de 5 ori până au arătat ca un telescop Cassegrain, nu ca o floare."),
    ("Î40. Ioan, ce ți-a fost cel mai greu de explicat lui Albert?",
     "(Ioan) Legea inversului pătratului aplicată la L4/L5. Distanța 150M km sună la fel de mare cu 7M km, dar intensitatea scade cu raportul pătratelor (~460×). Albert voia inițial să transmitem direct Soare→Pământ. A trebuit să-i arăt calculul: dacă laserul ar veni direct de la 7M km, la Pământ am avea <0,001% din putere. Relay-ul L4/L5 nu e opțional — e necesitate fizică."),
    ("Î41. Albert, când ai fost cel mai mândru de Ioan?",
     "(Albert) Când a observat singur că NbTi de la CERN nu e bun pentru baloane. Citea despre LHC și mi-a spus: „Albert, magneții ăștia trebuie răciți la 1,9 K cu heliu. Pentru baloanele noastre e imposibil.\" Apoi a căutat alternative și a găsit YBCO cu T_c 93 K. De acolo a pornit întregul design criogenic. Fără el, aveam cabluri supraconductoare… la temperatura camerei."),
    ("Î42. Câte ore pe săptămână dedicați proiectului?",
     "~6–8 ore în medie. Mai mult în weekend, mai puțin în săptămânile cu teze. În ultima lună înainte de județeană, ~15 ore/săptămână. Nu uităm școala — toate materiile contează pentru noi."),
    ("Î43. Ce materii din școală vă ajută cel mai mult?",
     "Fizica (evident — optică, radiații, termodinamică); Matematica (calcul de arii, exponențiale, trigonometrie); Informatica (cod pentru animații, simulări Three.js). Chimia ne ajută la înțelegerea YBCO (YBa₂Cu₃O₇). Geografia ne-a ajutat să alegem polii (latitudini, stratosferă). Practic, toate materiile — doar în proporții diferite."),
    ("Î44. Cum decideați direcția proiectului când apăreau alternative?",
     "Calcul obiectiv + argumentație. Nimic nu era „îmi place mai mult\". Ex.: 4 colectori vs 6 vs 8? Am calculat masa, costul lansării, eficiența colectării — 4 a rezultat optim. Nu „a simțit bine\", ci „numărul spune 4\"."),
    ("Î45. Ce rol a avut profesorul mentor?",
     "Prof. Francisc Aaron ne-a pus întrebări care dor. Nu a scris fișa, nu a făcut calculele. Dar când am zis „sistemul e stabil\", el a întrebat „dovediți că L4 rămâne stabil cu perturbații de 0,1%?\". Atunci am învățat. Mentorul = antrenor cu pretenții, nu co-autor."),
    ("Î46. Ce ați face diferit dacă ați lua proiectul de la zero?",
     "(a) Am începe cu simulare numerică din ziua 1, nu după 3 luni de calcule manuale. (b) Am aloca mai mult timp componentei L4/L5 (am subestimat complexitatea iris-ului). (c) Am face posterul în paralel cu calculele, nu la final — e un feedback vizual continuu."),
]

SET4 = [
    ("Î20. Dacă Soarele are o erupție solară masivă, P.O.L.A.R.I.S. se topește?",
     "Nu. Oglinzile au reflectivitate 90% — absorb doar 10% din energia incidentă și radiază termic în spațiu. Într-o erupție majoră (CME), fluxul crește de 2–3× pentru ore, nu zile. Sateliții colectori pot roti în fază inactivă (shutter) în <60 secunde. Sistemul e gradat pentru vârfuri — exact ca panourile SOHO sau Solar Dynamics Observatory."),
    ("Î21. Nu e P.O.L.A.R.I.S. o armă laser gigantică camuflată în proiect energetic?",
     "Tehnic, da — orice sistem de transmisie laser de 1,4 TW poate fi deturnat. De aceea arhitectura include două garanții: (1) razele au divergență mică dar focalizate pe iris-urile fixe L4/L5 — orice deviere face razele să nu mai aibă unde ateriza; (2) guvernanța propusă = operare multinațională ONU, similar ITER. Problema nu e tehnică, e politică — iar asta recunoaștem onest."),
    ("Î22. De ce nu panouri solare clasice pe Lună? Ar fi de 100× mai ieftin.",
     "Ar fi mai ieftin de construit pe Lună. Dar Luna are noapte de 14 zile — jumătate din timp e inactivă, factor de capacitate ~50%. În plus, transmisia Lună→Pământ are aceleași probleme ca a noastră, plus o complicație: praful lunar (regolit) degradează panourile rapid. P.O.L.A.R.I.S. la 7M km = 24/7 la 100%, fără atmosferă, fără praf."),
    ("Î23. Voi ați făcut proiectul sau profesorul vostru?",
     "Calculele, arhitectura, alegerea tehnologiilor, posterul, site-ul, codul pentru animație — noi. Profesorul Aaron ne-a verificat și ne-a pus întrebări incomode, care au făcut proiectul mai bun. Dar aportul e al nostru — și putem demonstra asta răspunzând la ORICE întrebare aveți, nu doar la cele repetate."),
    ("Î24. Ce se întâmplă dacă un meteorit lovește o oglindă?",
     "Probabilitate extrem de mică la 7M km (zona e „curățată\" de radiația solară de praful interplanetar). Dar dacă se întâmplă: pierdem 1 din 30 oglinzi = 3,3% reducere putere → 1,55 TW în loc de 1,6 TW. Redundanța e prin design. Înlocuirea dintr-un stoc orbital de rezervă: câteva săptămâni."),
    ("Î25. Cum răspundeți unui astrofizician care spune „proiectul e science-fiction\"?",
     "Fiecare componentă individuală există deja sau e demonstrată: Parker Solar Probe → navigație aproape de Soare ✓; JAXA → transmisie energie prin microunde ✓; LHC/CERN → supraconductori industriali ✓; Project Loon → baloane stratosferice ✓; oglinzi mari în spațiu → James Webb ✓. Nu inventăm fizică nouă. Integrăm tehnologii existente la scară mai mare. Asta nu e sci-fi, e inginerie."),
    ("Î26. Dacă ați avea 1 milion de euro mâine, ce ați face cu el pentru proiect?",
     "Un demonstrator la scară redusă — 1 oglindă Ø 10 m în LEO + 1 colector pe ISS + 1 laser de 1 kW → transmisie către un receptor în deșertul Gobi. Scara 1:100.000, dar dovada conceptului. Cost estimat: ~800.000 € hardware + 200.000 € launch via Falcon 9 rideshare."),
    ("Î27. De ce poartă echipa numele „POLARIS Bears\"? Urși în spațiu?",
     "Două motive: (1) POLARIS = steaua polară = referința la stațiile polare ale proiectului + „platforma\" care ne ghidează conceptual; (2) Bears = urșii polari trăiesc la poli, acolo unde ajung laserele noastre. Sunt simbol al rezistenței în medii extreme. Și mai sincer: amândoi iubim trenurile și ursul din stema familiilor noastre. E un nume personal, nu de marketing."),
    ("Î47. Un AI avansat ar înlocui sistemul vostru cu ceva mai simplu. Ce?",
     "Probabil nano-botanie — bacterii modificate genetic care produc hidrogen din lumina solară la ~50% eficiență (teoretic). Sau fuziune aneutronică (He-3 de pe Lună). P.O.L.A.R.I.S. e soluția pentru tehnologia actuală extrapolată. Un AI cu descoperiri noi de fizică ar găsi ceva mai elegant. Dar azi, P.O.L.A.R.I.S. e cel mai bun drum."),
    ("Î48. Ce s-ar întâmpla dacă laserul ratează iris-ul și lovește Stația Spațială Internațională?",
     "ISS orbitează la ~400 km altitudine. L4/L5 e la 150M km. Laser 2,83 × 10¹² W concentrat pe Ø ~1 km la țintă. Dacă ratează și lovește ISS, intensitatea ~3,6 MW/m² — ar topi ISS-ul în secunde. De aceea: (a) sistemele de siguranță ale iris-ului fac cut-off automat în <100 ms la pierderea țintei; (b) geometria orbitei exclude ISS ca obstacol permanent. Calculul risc vs beneficiu e crucial și parte din designul de siguranță."),
    ("Î49. De ce în română și nu în engleză (pentru EUCYS)?",
     "ONCS este competiție românească. Prezentarea pentru juriu trebuie în română — e regulamentul. Pentru EUCYS (dacă ne calificăm), avem versiune engleză pregătită: fișa, posterul, slide-urile. Nu suntem limitați lingvistic — suntem adaptați contextului."),
    ("Î50. Dacă Rusia/China/SUA n-ar coopera, proiectul moare?",
     "Moare la scară globală, da. Dar ideile nu mor. Ne uităm la ISS — cooperare SUA-Rusia de 25 de ani, în plin război rece 2.0. Proiectele spațiale forțează cooperarea pentru că nimeni nu poate singur. Dacă P.O.L.A.R.I.S. n-ar veni, tensiunile geopolitice ar fi probabil mai mari, nu mai mici."),
    ("Î51. Care e lecția etică a P.O.L.A.R.I.S.?",
     "Energia nu trebuie să fie privilegiu. Azi, 770 milioane de oameni nu au electricitate. P.O.L.A.R.I.S. livrează spre poli = spre zone slab electrificate. Lecția: tehnologia megascale trebuie proiectată pentru egalitate, nu pentru profit. De aceea surplusul (53%) merge la Star Power Grid pentru zone izolate, nu la licitație privată."),
    ("Î52. Dacă ar trebui să tăiați 30% din bugetul proiectului, ce scoateți?",
     "Am scoate Star Power Grid (extensia viitoare). E valoroasă, dar nu esențială pentru primul 1,6 TW. Dyson Ring + Colectori + L4/L5 + centrale polare = MVP (Minimum Viable Polaris). Star Power Grid = faza 2."),
    ("Î53. Ce cărți sau resurse ați citit pe subiect?",
     "Freeman Dyson — „Search for Artificial Stellar Sources\" (1960, lucrarea originală); Carl Sagan — „Contact\" (pentru scale civilizaționale); NASA Technical Reports — „Solar Power Satellites\" (Glaser, 1978); OpenStax Physics (electromagnetism, supraconductivitate); Isaac Arthur pe YouTube (canalul „Science & Futurism\", serie întreagă despre Dyson Swarms). Nu doar manuale — inclusiv science-fiction rigoros (Arthur C. Clarke)."),
    ("Î54. Peste 20 de ani, ce veți spune că a însemnat proiectul acesta?",
     "(Ioan) „Proiectul care m-a învățat că fizica are răspunsuri doar dacă pui întrebarea corectă. Și m-a învățat să calculez fără frică de numere mari.\"\n(Albert) „Proiectul care m-a făcut să înțeleg că o idee frumoasă nu trebuie doar să fie corectă — trebuie să fie comunicată. Jumătate din muncă e în matematică, cealaltă jumătate în cum spui matematica aceea.\""),
]

# ----------------------------- STYLING HELPERS -----------------------------

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(6)  # WD_BREAK.PAGE = 7 actually

def para(doc, text, *, bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

# ----------------------------- DOCUMENT BUILD -----------------------------

doc = Document()

# page margins
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ----------------- COVER -----------------
para(doc, "P.O.L.A.R.I.S.", bold=True, size=28, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "54 Întrebări & Răspunsuri pentru Juriul ONCS 2026",
     bold=True, size=16, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "Echipa POLARIS Bears — Ioan Cristian CHELARU & Albert David OLARIU",
     italic=True, size=11, color=CYAN, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "Secțiunea C — Tehnologia Informației · Juniori · Clasa a VIII-a",
     italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "Mentor: Prof. Francisc Dionisie Aaron",
     italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

# intro
intro = (
    "Acest document conține 54 de întrebări pe care juriul ONCS 2026 le-ar putea adresa, "
    "împreună cu răspunsurile recomandate. Întrebările sunt organizate în 4 seturi tematice, "
    "alinate la criteriile de evaluare din Regulamentul ONCS (Art. 16 — Evaluarea) și la "
    "punctele de interes ale juriului (Anexa 1, secțiunea c — Posibil Chestionar pentru "
    "Participanți). La final găsiți un tabel cu sfaturi practice pentru administrarea răspunsurilor "
    "în timpul prezentării."
)
para(doc, intro, size=11, space_after=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# set mapping table
para(doc, "Harta seturilor → criteriile de punctaj", bold=True, size=12, color=NAVY, space_after=4)

mapping_data = [
    ("Set", "Temă", "Criterii acoperite", "Punctaj maxim influențat"),
    ("1", "Științific & Tehnic", "(b) abordare științifică, (d) complexitatea cercetării", "40 p"),
    ("2", "Viziune, Problemă, Impact", "(a) creativitate, (c) obiective", "40 p"),
    ("3", "Echipă & Proces", "(f) activitatea elevilor + muncă în echipă", "10 p (transversal)"),
    ("4", "Atipic — gândire laterală", "(a), (b), (d), (e) — evaluat în Q&A", "variabil"),
]
table = doc.add_table(rows=len(mapping_data), cols=4)
table.autofit = False
widths = [Cm(1.2), Cm(4.5), Cm(7.0), Cm(4.3)]
for i, row in enumerate(mapping_data):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.width = widths[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, '0F131F')
        else:
            set_cell_shading(cell, 'F5F5F5' if i % 2 else 'FFFFFF')

para(doc, "", size=8, space_after=4)
doc.add_page_break()

# ----------------- SETS -----------------

def render_set(doc, title_text, subtitle, color, qa_list):
    para(doc, title_text, bold=True, size=18, color=color, space_after=2)
    para(doc, subtitle, italic=True, size=10, color=RGBColor(0x55, 0x55, 0x55), space_after=10)

    for q, a in qa_list:
        # question
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(q)
        r.bold = True
        r.font.size = Pt(11.5)
        r.font.color.rgb = NAVY

        # answer (handle newlines as sub-paragraphs)
        for line in a.split('\n'):
            pa = doc.add_paragraph()
            pa.paragraph_format.space_after = Pt(8)
            pa.paragraph_format.left_indent = Cm(0.5)
            pa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            ra = pa.add_run("R: " if line == a.split('\n')[0] else "    ")
            ra.bold = True
            ra.font.size = Pt(10.5)
            ra.font.color.rgb = GOLD
            rb = pa.add_run(line)
            rb.font.size = Pt(10.5)

render_set(doc,
           "SET 1 — ȘTIINȚIFIC & TEHNIC",
           "Criteriile (b) 20p + (d) 20p. Răspunde prioritar Ioan; Albert completează cu inginerie/materiale.",
           NAVY, SET1)
doc.add_page_break()

render_set(doc,
           "SET 2 — VIZIUNE, PROBLEMĂ, IMPACT",
           "Criteriile (a) 30p + (c) 10p. Răspunde prioritar Albert; Ioan completează cu cifre de comparație.",
           NAVY, SET2)
doc.add_page_break()

render_set(doc,
           "SET 3 — ECHIPĂ & PROCES",
           "Criteriul (f) 10p (transversal). CRITIC: juriul testează fiecare elev separat. Fiecare răspunde și despre sine, și despre celălalt.",
           NAVY, SET3)
doc.add_page_break()

render_set(doc,
           "SET 4 — ATIPIC (Întrebări-capcană)",
           "Întrebări imprevizibile din partea juriului, care testează gândirea laterală, onestitatea și profunzimea.",
           NAVY, SET4)

# ----------------- SFAT PRACTIC -----------------
doc.add_page_break()
para(doc, "Cum administrați cele 4 seturi — sfat practic",
     bold=True, size=16, color=GOLD, space_after=8)

advice = [
    ("Set", "Cine răspunde primul", "Timp mediu", "Material vizibil"),
    ("1 — Științific", "Ioan → Albert completează", "60–90 sec", "Poster (formule, lanț energetic)"),
    ("2 — Viziune", "Albert → Ioan completează", "45–60 sec", "PPT (diagrame Kardashev)"),
    ("3 — Echipă", "Individual — fiecare separat", "30–45 sec", "Timeline proiect pe poster"),
    ("4 — Atipic", "Cel mai apropiat de temă", "45–60 sec", "Fără suport — gândire pe loc"),
]
table = doc.add_table(rows=len(advice), cols=4)
table.autofit = False
widths = [Cm(3.0), Cm(5.0), Cm(2.5), Cm(6.5)]
for i, row in enumerate(advice):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.width = widths[j]
        cell.text = ''
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        if i == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, '0F131F')
        else:
            set_cell_shading(cell, 'F5F5F5' if i % 2 else 'FFFFFF')

para(doc, "", size=8, space_after=6)

para(doc, "Regula de aur",
     bold=True, size=13, color=GOLD, space_after=4)
golden_rule = (
    "Dacă nu știți răspunsul la o întrebare atipică, spuneți onest: "
    "„E o întrebare la care nu am gândit în această formă. Reflexia noastră imediată este X, "
    "dar e ceva ce vrem să aprofundăm.\" Juriul ONCS apreciază onestitatea intelectuală "
    "mai mult decât improvizația falsă. A nu ști ≠ a fi nepregătit — A nu recunoaște că nu știi = a fi nepregătit."
)
para(doc, golden_rule, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=14)

# ----------------- CHECKLIST -----------------
para(doc, "Checklist înainte de prezentare",
     bold=True, size=13, color=GOLD, space_after=4)
checklist = [
    "Fiecare a citit de cel puțin 3 ori ambele seturi (1+2) — trebuie să poată răspunde la ORICE, indiferent de „specializare\".",
    "Setul 3 (Echipă) a fost exersat INDIVIDUAL — Ioan răspunde la Î13 fără Albert în cameră; Albert la Î14 fără Ioan.",
    "Setul 4 (Atipic) a fost parcurs cel puțin o dată împreună, discutând variante de răspuns.",
    "Toate numerele cheie (1,6 TW, 48%, 7M km, 30 oglinzi, 93 K, 800 GW) sunt memorate, nu citite.",
    "Pregătit un răspuns de rezervă pentru: „Nu știu. Am o intuiție, dar nu am calcul exact.\"",
    "Discutat în echipă: cum nu vă contraziceți reciproc (dacă nu sunteți siguri, mai bine tăceți decât să vă corectați).",
]
for item in checklist:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(item)
    r.font.size = Pt(10.5)

# ----------------- FOOTER -----------------
para(doc, "", size=8, space_after=4)
footer_text = (
    "Document generat pe baza Regulamentului specific ONCS 2026 "
    "(nr. 24.664 / 04.02.2026) și a datelor științifice canonice ale proiectului P.O.L.A.R.I.S. "
    "© 2026 Echipa POLARIS Bears — polaris-bears.ro"
)
para(doc, footer_text, italic=True, size=8, color=RGBColor(0x88, 0x88, 0x88),
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

doc.save(OUTPUT)
print(f"OK: salvat la {OUTPUT}")
print(f"Total intrebari: {len(SET1) + len(SET2) + len(SET3) + len(SET4)}")

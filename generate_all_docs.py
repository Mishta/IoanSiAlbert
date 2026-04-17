#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_all_docs.py  --  ONCS 2026 / P.O.L.A.R.I.S.

Genereaza 5 documente DOCX separate in output/:
  01_Coperta.docx
  02_Cuprins.docx
  03_Anexa_2_Date_Identificare.docx
  04_Rezumat.docx
  05_Document_Principal.docx   (max 10 pagini, cu imagini)
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE  = "c:/Antigravity/__SKILLS/IoanSiAlbert"
OUT   = os.path.join(BASE, "output")
IMG   = os.path.join(BASE, "images", "generated")
DIAG  = os.path.join(BASE, "assets", "diagrams")
ASS   = os.path.join(BASE, "assets")

# ─── Helpers ──────────────────────────────────────────────────────────────────

def mkdoc():
    doc = Document()
    for s in doc.sections:
        s.top_margin    = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin   = Cm(2.5)
        s.right_margin  = Cm(2.5)
    n = doc.styles['Normal']
    n.font.name = 'Calibri'
    n.font.size = Pt(10.5)
    return doc

def hd(doc, txt, lvl=1, sz=None, sb=6, sa=3):
    ph = doc.add_heading(txt, level=lvl)
    ph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ph.paragraph_format.space_before = Pt(sb)
    ph.paragraph_format.space_after  = Pt(sa)
    if sz:
        for r in ph.runs:
            r.font.size = Pt(sz)
    return ph

def tx(doc, txt, bold=False, italic=False, sz=10.5, sa=3,
        al=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pp = doc.add_paragraph()
    pp.alignment = al
    pp.paragraph_format.space_after = Pt(sa)
    rr = pp.add_run(txt)
    rr.bold  = bold
    rr.italic = italic
    rr.font.size = Pt(sz)
    rr.font.name = 'Calibri'
    return pp

def blt(doc, txt, sz=10, sa=1):
    pp = doc.add_paragraph(style='List Bullet')
    pp.paragraph_format.space_after = Pt(sa)
    rr = pp.add_run(txt)
    rr.font.size = Pt(sz)
    rr.font.name = 'Calibri'
    return pp

def pic(doc, path, w=8.0, cap=None):
    try:
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(2)
        pp.paragraph_format.space_after  = Pt(1)
        pp.add_run().add_picture(path, width=Cm(w))
    except Exception:
        tx(doc, "[Imagine: " + os.path.basename(path) + "]",
           italic=True, al=WD_ALIGN_PARAGRAPH.CENTER)
    if cap:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(4)
        rr = cp.add_run(cap)
        rr.italic = True
        rr.font.size = Pt(8.5)
        rr.font.name = 'Calibri'

def set_cw(cell, cm):
    tc  = cell._tc
    tp  = tc.get_or_add_tcPr()
    tw  = OxmlElement('w:tcW')
    tw.set(qn('w:w'), str(int(cm * 567)))
    tw.set(qn('w:type'), 'dxa')
    tp.append(tw)

def cell_tx(cell, txt, bold=False, italic=False, sz=10.5,
            al=WD_ALIGN_PARAGRAPH.LEFT):
    pp = cell.paragraphs[0]
    pp.alignment = al
    pp.paragraph_format.space_after = Pt(2)
    rr = pp.add_run(txt)
    rr.bold   = bold
    rr.italic = italic
    rr.font.size = Pt(sz)
    rr.font.name = 'Calibri'

# ─── DOC 1: COPERTA ───────────────────────────────────────────────────────────

def doc1_coperta():
    doc = mkdoc()
    for _ in range(3):
        doc.add_paragraph()

    logo = os.path.join(ASS, "logo-no-ONCS-text-bright-highrez.png")
    pic(doc, logo, w=5.5)
    doc.add_paragraph()

    tx(doc, "P.O.L.A.R.I.S.",
       bold=True, sz=26, al=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
    tx(doc, "Platform\u0103 Orbital\u0103 Laser pentru",
       sz=13, sa=1, al=WD_ALIGN_PARAGRAPH.CENTER)
    tx(doc, "Alimentare \u015fi Recep\u021bie Inovativ\u0103 Solar\u0103",
       sz=13, sa=6, al=WD_ALIGN_PARAGRAPH.CENTER)
    tx(doc, "Proiect \u015ftiin\u021bific",
       italic=True, sz=11, sa=2, al=WD_ALIGN_PARAGRAPH.CENTER)
    tx(doc, "Analiz\u0103 fizic\u0103 \u2022 Inginerie spa\u021bial\u0103 \u2022 Infrastructur\u0103 energetic\u0103 global\u0103",
       italic=True, sz=10, sa=10, al=WD_ALIGN_PARAGRAPH.CENTER)

    tx(doc, "Olimpiada Na\u021bional\u0103 de Creativitate \u0218tiin\u021bific\u0103 \u2014 ONCS 2026",
       bold=True, sz=12, sa=2, al=WD_ALIGN_PARAGRAPH.CENTER)
    tx(doc, "Sec\u021biunea: \u0218tiin\u021be fundamentale",
       sz=11, sa=10, al=WD_ALIGN_PARAGRAPH.CENTER)

    tx(doc, "Echip\u0103: POLARIS Bears",
       bold=True, sz=13, sa=4, al=WD_ALIGN_PARAGRAPH.CENTER)
    tx(doc, "Ioan Cristian CHELARU | Albert David OLARIU",
       sz=11, sa=2, al=WD_ALIGN_PARAGRAPH.CENTER)
    tx(doc, "clasa a VIII-a, ICHB \u2014 Liceul Teoretic Interna\u021bional de Informatic\u0103 Bucure\u015fti",
       sz=10, sa=6, al=WD_ALIGN_PARAGRAPH.CENTER)
    tx(doc, "Mentor: Prof. Francisc Dionisie AARON",
       sz=11, sa=1, al=WD_ALIGN_PARAGRAPH.CENTER)
    tx(doc, "ICHB \u2014 Liceul Teoretic Interna\u021bional de Informatic\u0103",
       sz=10, sa=6, al=WD_ALIGN_PARAGRAPH.CENTER)

    tx(doc, "Colaboratori:", bold=True, sz=11, sa=2,
       al=WD_ALIGN_PARAGRAPH.CENTER)
    tx(doc, "Prof. Univ. Dr. Valentin BARNA \u2014 Facultatea de Fizic\u0103, Univ. Bucure\u015fti",
       sz=10, sa=1, al=WD_ALIGN_PARAGRAPH.CENTER)
    tx(doc, "Carlo Emilio MONTANARI \u2014 Fizician aplicat, Doctorand CERN / Univ. Bologna",
       sz=10, sa=1, al=WD_ALIGN_PARAGRAPH.CENTER)
    tx(doc, "Daniel-Justinian ZELENSCHI \u2014 NASA Ames Research Center",
       sz=10, sa=8, al=WD_ALIGN_PARAGRAPH.CENTER)

    tx(doc, "Aprilie 2026",
       sz=11, bold=True, al=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    tx(doc,
       "Toate datele de identificare sunt completate. "
       "Verificati diacriticele si formatarea inainte de depunerea dosarului.",
       italic=True, sz=9, al=WD_ALIGN_PARAGRAPH.CENTER, sa=2)

    path = os.path.join(OUT, "01_Coperta.docx")
    doc.save(path)
    print("  [OK] 01_Coperta.docx")

# ─── DOC 2: CUPRINS ───────────────────────────────────────────────────────────

def doc2_cuprins():
    doc = mkdoc()
    hd(doc, "CUPRINS", lvl=1, sz=16, sb=4, sa=10)
    tx(doc, "Document de referin\u021b\u0103: 05_Document_Principal.docx",
       italic=True, sz=9, sa=8)

    entries = [
        ("1.",       "Introducere",                                                  "1",  False),
        ("2.",       "Fundamente Teoretice",                                         "1",  False),
        ("  2.1",    "Concepte de baz\u0103",                                        "1",  True),
        ("  2.2",    "Legi, principii \u015fi teorii aplicabile",                    "2",  True),
        ("  2.3",    "Modele matematice relevante",                                  "2",  True),
        ("3.",       "Originalitatea \u015fi Inova\u021bia Proiectului",             "3",  False),
        ("4.",       "Obiectivele Cercet\u0103rii",                                  "3",  False),
        ("  4.1",    "Obiectiv General",                                             "3",  True),
        ("  4.2",    "Obiective Specifice (m\u0103surabile)",                        "3",  True),
        ("5.",       "Ipoteze \u015fi Predic\u021bii",                               "4",  False),
        ("6.",       "Materiale \u015fi Metode",                                     "4",  False),
        ("  6.1",    "Materiale",                                                    "4",  True),
        ("  6.2",    "Metodologie",                                                  "4",  True),
        ("7.",       "Arhitectura Sistemului Energetic",                             "5",  False),
        ("8.",       "STAR POWER GRID \u2014 Sintez\u0103",                          "6",  False),
        ("9.",       "Rezultate \u2014 Fluxul Energetic Complet",                    "6",  False),
        ("  9.1",    "Date brute (input)",                                           "6",  True),
        ("  9.2\u20139.7", "Pa\u015fii de calcul \u015fi sinteza rezultatelor",     "7",  True),
        ("10.",      "Estimare de Costuri",                                          "8",  False),
        ("11.",      "Analiza Rezultatelor",                                         "8",  False),
        ("  11.1",   "Fluxul energetic de la Soare la L4/L5",                       "8",  True),
        ("  11.2",   "Limitarea la 2 \u00d7 800 GW prin YBCO",                      "8",  True),
        ("  11.3",   "Masa total\u0103 \u015fi flotabilitatea baloanelor",           "9",  True),
        ("  11.4",   "Impactul energetic global (1,6 TW)",                          "9",  True),
        ("  11.5",   "Star Power Grid \u2014 analiz\u0103 conceptual\u0103",         "9",  True),
        ("  11.6",   "Tranzi\u021bia pe Scala Kardashev",                            "9",  True),
        ("12.",      "Discu\u021bii, Limit\u0103ri \u015fi Concluzii",               "10", False),
        ("",         "Bibliografie",                                                 "10", False),
    ]

    tbl = doc.add_table(rows=len(entries), cols=3)
    tbl.style = 'Normal Table'
    for i, (num, title, pg, is_sub) in enumerate(entries):
        row = tbl.rows[i]
        sz  = 9.5 if is_sub else 10.5

        # col 0 – numar
        c0 = row.cells[0]
        set_cw(c0, 1.2)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(num)
        r0.font.size = Pt(sz); r0.font.name = 'Calibri'
        r0.bold = not is_sub and num != ""

        # col 1 – titlu
        c1 = row.cells[1]
        set_cw(c1, 12.0)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(title)
        r1.font.size = Pt(sz); r1.font.name = 'Calibri'
        r1.bold = not is_sub and num != ""

        # col 2 – pagina
        c2 = row.cells[2]
        set_cw(c2, 1.5)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run("p. " + pg)
        r2.font.size = Pt(sz); r2.font.name = 'Calibri'

    path = os.path.join(OUT, "02_Cuprins.docx")
    doc.save(path)
    print("  [OK] 02_Cuprins.docx")

# ─── DOC 3: DATE IDENTIFICARE ─────────────────────────────────────────────────

def doc3_identificare():
    doc = mkdoc()
    hd(doc, "Date de Identificare \u2014 Dosar \u00cens\u0063riere ONCS 2026",
       lvl=1, sz=14, sb=4, sa=6)
    tx(doc,
       "Anex\u0103 nr. 2 \u2014 Regulament ONCS 2026 | "
       "\u26a0  C\u00e2mpurile marcate necesit\u0103 completare \u00eenainte de depunere.",
       italic=True, sz=9, sa=8)

    rows = [
        ("Acronimul proiectului",   "P.O.L.A.R.I.S."),
        ("Titlul complet",
         "Platform\u0103 Orbital\u0103 Laser pentru Alimentare \u015fi Recep\u021bie Inovativ\u0103 Solar\u0103"),
        ("Categoria ONCS 2026",     "Juniori"),
        ("Sec\u021biunea",          "\u0218tiin\u021be fundamentale"),
        ("Mentorul echipei",
         "Prof. Francisc Dionisie AARON\n"
         "ICHB \u2014 Liceul Teoretic Interna\u021bional de Informatic\u0103 Bucure\u015fti\n"
         "Nr. telefon: 0721 185 014"),
        ("Elev 1",
         "Ioan Cristian CHELARU\n"
         "Clasa a VIII-a | ICHB \u2014 Liceul Teoretic Interna\u021bional de Informatic\u0103 Bucure\u015fti\n"
         "Rol: Modelare vizual\u0103, scala Kardashev, re\u021beaua Star Power Grid, prezentare"),
        ("Elev 2",
         "Albert David OLARIU\n"
         "Clasa a VIII-a | ICHB \u2014 Liceul Teoretic Interna\u021bional de Informatic\u0103 Bucure\u015fti\n"
         "Rol: Arhitectur\u0103 sistem, calcule energetice, supraconductori YBCO, dinamic\u0103 orbital\u0103"),
        ("Colaborator 1",
         "Prof. Univ. Dr. Valentin BARNA\n"
         "Facultatea de Fizic\u0103, Univ. Bucure\u015fti\n"
         "Specializare: Optic\u0103 \u015fi Fotonic\u0103, Fizica LASER-ilor\n"
         "Consultan\u021b\u0103: principii de fotonic\u0103 \u015fi conversie radia\u021bie solar\u0103 \u2192 LASER"),
        ("Colaborator 2",
         "Carlo Emilio MONTANARI\n"
         "Fizician aplicat \u2022 Doctorand Dinamica Fasciculelor, CERN / Univ. Bologna\n"
         "Consultan\u021b\u0103: dinamica fasciculelor de \u00eenalt\u0103 precizie (LHC \u2192 POLARIS)"),
        ("Colaborator 3",
         "Daniel-Justinian ZELENSCHI\n"
         "Ing. chimic (UCL + Melbourne) \u2022 PhD(c) AI \u2022 NASA Ames Research Center\n"
         "Space Settlement Contest \u2014 Premiul II \u2022 implementare AI la scar\u0103 industrial\u0103 (Raiffeisen AI)\n"
         "Consultan\u021b\u0103: algoritmi AI pentru controlul autonom al swarm-ului orbital \u015fi optimizarea Star Power Grid"),
        ("Coordonator \u015ftiin\u021bific",
         "Prof. Francisc Dionisie AARON\n"
         "ICHB \u2014 Liceul Teoretic Interna\u021bional de Informatic\u0103 Bucure\u015fti"),
    ]

    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    for i, (label, val) in enumerate(rows):
        row = tbl.rows[i]
        c0  = row.cells[0]
        c1  = row.cells[1]
        set_cw(c0, 4.5)
        set_cw(c1, 11.0)

        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(3)
        r0 = p0.add_run(label)
        r0.bold = True; r0.font.size = Pt(10); r0.font.name = 'Calibri'

        lines = val.split('\n')
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(1)
        r1 = p1.add_run(lines[0])
        r1.font.size = Pt(10); r1.font.name = 'Calibri'
        for line in lines[1:]:
            pp = c1.add_paragraph()
            pp.paragraph_format.space_after = Pt(1)
            rr = pp.add_run(line)
            rr.font.size = Pt(9.5); rr.font.name = 'Calibri'
            rr.italic = line.startswith('\u26a0')

    doc.add_paragraph()
    tx(doc,
       "Not\u0103: Completarea datelor se va face cu denumiri corecte, cu diacritice, "
       "f\u0103r\u0103 prescurt\u0103ri. Documentul se va transmite \u00een format editabil \u015fi scanat.",
       italic=True, sz=9, sa=2)

    path = os.path.join(OUT, "03_Anexa_2_Date_Identificare.docx")
    doc.save(path)
    print("  [OK] 03_Anexa_2_Date_Identificare.docx")

# ─── DOC 4: REZUMAT ───────────────────────────────────────────────────────────

def doc4_rezumat():
    doc = mkdoc()
    hd(doc, "Rezumat \u2014 P.O.L.A.R.I.S.", lvl=1, sz=14, sb=4, sa=6)
    tx(doc, "Olimpiada Na\u021bional\u0103 de Creativitate \u0218tiin\u021bific\u0103 \u2014 ONCS 2026 | "
       "Maxim 250 de cuvinte",
       italic=True, sz=9, sa=8)

    tx(doc,
       "Criza energetic\u0103 global\u0103 \u015fi epuizarea treptat\u0103 a combustibililor fosili "
       "reprezint\u0103 provocarea secolului XXI. P.O.L.A.R.I.S. "
       "(Platform\u0103 Orbital\u0103 Laser pentru Alimentare \u015fi Recep\u021bie Inovativ\u0103 Solar\u0103) "
       "propune o arhitectur\u0103 energetic\u0103 spa\u021bial\u0103 complet\u0103: captarea energiei solare "
       "de la 7 milioane km fa\u021b\u0103 de Soare, conversia \u00een fascicule LASER "
       "\u015fi transmiterea spre P\u0103m\u00e2nt.")

    tx(doc,
       "Sistemul utilizeaz\u0103 un Dyson swarm de 30 de oglinzi orbitale de 1 km diametru fiecare, "
       "plasate la aproximativ 7 milioane km de Soare, care reflect\u0103 radia\u021bia solar\u0103 "
       "c\u0103tre 4 satelit\u021bi colectori. Ace\u015ftia transform\u0103 lumina \u00een fascicule LASER, "
       "direc\u021bionate c\u0103tre noduri energetice amplasate \u00een punctele Lagrange L4 \u015fi L5. "
       "Centralele orbitale polare, sus\u021binute la altitudini de ~30 km prin baloane cu heliu, "
       "convertesc energia LASER \u00een electricitate \u015fi o transmit la sol prin cabluri "
       "supraconductoare, fiecare capabil de 800 GW, rezultând o putere total\u0103 de ~1,6 TW. "
       "Livra\u021bi continuu reprezint\u0103 \u224848% din consumul electric global al anului 2023. "
       "Cablurile supraconductoare YBCO, validate prin consultan\u021b\u0103 cu fizicieni de la CERN, "
       "opereaz\u0103 la 77 K cu azot lichid, transmit\u00e2nd 800 GW per cablu "
       "la o mas\u0103 de ~7 t per segment de 30 km.")

    tx(doc,
       "Extensia Star Power Grid distribuie energia radiant de la ambii poli spre continente, "
       "printr-o re\u021bea terestr\u0103 de \u00eenalt\u0103 tensiune. Prin aceast\u0103 arhitectur\u0103, "
       "umanitatea ar avansa de la Tipul 0,73 spre Tipul 1,0 pe scala Kardashev, "
       "prima civiliza\u021bie care st\u0103p\u00e2ne\u015fte integral energia propriei planete.")

    tx(doc,
       "De\u015fi costurile dep\u0103\u015fesc 32.000 miliarde EUR la tehnologia actual\u0103, "
       "proiectul integreaz\u0103 fizica clasic\u0103, supraconductivitatea, criogenia "
       "\u015fi ingineria orbital\u0103 \u00eentr-un model teoretic riguros, "
       "cu relevan\u021b\u0103 educa\u021bional\u0103 \u015fi \u015ftiin\u021bific\u0103.",
       sa=8)

    tx(doc, "Num\u0103r estimat de cuvinte: ~249", italic=True, sz=9,
       al=WD_ALIGN_PARAGRAPH.RIGHT)

    path = os.path.join(OUT, "04_Rezumat.docx")
    doc.save(path)
    print("  [OK] 04_Rezumat.docx")

# ─── DOC 5: DOCUMENT PRINCIPAL (max 10 pag.) ──────────────────────────────────

def doc5_principal():
    doc = mkdoc()

    # ── Pagina 1: Titlu + Cap. 1 ─────────────────────────────────────────────
    logo = os.path.join(ASS, "logo-no-ONCS-text-bright-highrez.png")
    pic(doc, logo, w=3.5)
    tx(doc, "P.O.L.A.R.I.S.",
       bold=True, sz=20, al=WD_ALIGN_PARAGRAPH.CENTER, sa=1)
    tx(doc, "Platform\u0103 Orbital\u0103 Laser pentru Alimentare \u015fi Recep\u021bie Inovativ\u0103 Solar\u0103",
       italic=True, sz=10, al=WD_ALIGN_PARAGRAPH.CENTER, sa=1)
    tx(doc, "POLARIS Bears \u2022 ONCS 2026 \u2022 Sec\u021biunea: \u0218tiin\u021be fundamentale",
       sz=9, al=WD_ALIGN_PARAGRAPH.CENTER, sa=6)

    hd(doc, "1. Introducere", lvl=1, sz=12, sb=4, sa=2)
    tx(doc,
       "Consumul global de energie este \u00een cre\u015ftere continu\u0103, amplific\u00e2nd dependen\u021ba "
       "de combustibili fosili \u015fi riscurile de mediu aferente. Energia solar\u0103 terestr\u0103, "
       "de\u015fi abundent\u0103, este limitat\u0103 de absorb\u021bia atmosferic\u0103 (~30%), de ciclul "
       "zi-noapte \u015fi de sezonalitate. Un sistem de captare spa\u021bial\u0103, \u00een proximitatea Soarelui, "
       "dep\u0103\u015fe\u015fte toate aceste limit\u0103ri: intensitatea radia\u021biei cre\u015fte de la "
       "1.361 W/m\u00b2 (la 1 UA) la \u22486,22 \u00d7 10\u2075 W/m\u00b2 la 7 milioane km, iar "
       "disponibilitatea devine practic continu\u0103.")
    tx(doc,
       "Obiectivul proiectului: elaborarea modelului teoretic al unui Sistem Energetic "
       "Spa\u021bial (SES) format din roi Dyson de oglinzi orbitale, satelit\u021bi colectori laser, "
       "noduri de transfer la punctele Lagrange L4 \u015fi L5 \u015fi sta\u021bii orbitale polare conectate "
       "la P\u0103m\u00e2nt prin cabluri supraconductoare YBCO \u2014 capabil s\u0103 livreze 2 \u00d7 800 GW.")
    tx(doc,
       "\u00centrebarea de cercetare: Poate un astfel de sistem s\u0103 constituie o solu\u021bie teoretic "
       "viabil\u0103 pentru producerea \u015fi transmiterea energiei la scar\u0103 global\u0103, "
       "\u00een limitele fizicii cunoscute?")

    pic(doc, os.path.join(IMG, "Image A - Earth Energy Crisis.png"), w=7.0,
        cap="Fig. 1 \u2014 Criza energetic\u0103 global\u0103: context de plecare al proiectului P.O.L.A.R.I.S.")

    # ── Cap. 2 ────────────────────────────────────────────────────────────────
    hd(doc, "2. Fundamente Teoretice", lvl=1, sz=12, sb=6, sa=2)
    hd(doc, "2.1 Concepte de baz\u0103", lvl=2, sz=11, sb=3, sa=2)
    tx(doc,
       "Radia\u021bia solar\u0103 \u015fi legea inversului p\u0103tratului: I(r) = I_Earth \u00d7 (d_SE / r)\u00b2. "
       "La r = 7 \u00d7 10\u2076 km, intensitatea ajunge la \u22486,22 \u00d7 10\u2075 W/m\u00b2 \u2014 "
       "de ~457 de ori mai mare dec\u00e2t la suprafa\u021ba P\u0103m\u00e2ntului.")
    tx(doc,
       "Punctele Lagrange L4 \u015fi L5: pozi\u021bii de echilibru gravita\u021bional stabil \u00een sistemul "
       "Soare\u2013P\u0103m\u00e2nt, ideale pentru amplasarea nodurilor energetice cu consum minim de combustibil.")
    tx(doc,
       "Supraconductori YBCO (YBa\u2082Cu\u2083O\u2087\u208b\u03b4): T_c \u2248 92 K; operare practic\u0103 "
       "la 77 K cu azot lichid (nu heliu la 4 K, ca NbTi clasic). Densitate de curent critic\u0103 j_c "
       "net superioar\u0103 NbTi la temperaturi ridicate; mas\u0103 specific\u0103 mult mai mic\u0103. "
       "Alegerea YBCO a fost validat\u0103 prin consultan\u021b\u0103 cu cercet\u0103tori CERN.")
    tx(doc,
       "Flotabilitate stratrosferic\u0103: F_A = \u03c1_aer \u00d7 V \u00d7 g. "
       "La 30 km, \u03c1_aer \u2248 0,018 kg/m\u00b3; un ansamblu de baloane cu heliu "
       "de c\u00e2\u021biva mii de m\u00b3 poate sus\u021bine ~11 t.")

    hd(doc, "2.2 Legi \u015fi principii aplicabile", lvl=2, sz=11, sb=3, sa=2)
    tx(doc,
       "Conservarea energiei: E_util = \u03b7 \u00d7 E_input, 0 < \u03b7 < 1. "
       "Eficien\u021be utilizate: reflectivitate oglinzi 90%, conversie lumin\u0103\u2192laser 43%, "
       "conversie laser\u2192electricitate 60%. "
       "Densitatea critic\u0103 de curent YBCO: j = I/A \u2264 j_c \u2014 "
       "cablul se dimensioneaz\u0103 astfel \u00eenc\u00e2t curentul s\u0103 nu dep\u0103\u015feasc\u0103 "
       "limita critic\u0103 a materialului.")

    hd(doc, "2.3 Modele matematice relevante", lvl=2, sz=11, sb=3, sa=2)
    tx(doc,
       "P_captat\u0103 = I_sol \u00d7 A_tot = 6,22\u00d710\u2075 \u00d7 2,36\u00d710\u2077 "
       "\u2248 1,46\u00d710\u00b9\u00b3 W     |     "
       "P_reflectat\u0103 = P_captat\u0103 \u00d7 0,90 \u2248 1,32\u00d710\u00b9\u00b3 W")
    tx(doc,
       "P_laser/sat = (P_reflectat\u0103 / 4) \u00d7 0,43 \u2248 1,42\u00d710\u00b9\u00b2 W     |     "
       "P \u00een L4 (sau L5) = 2 \u00d7 P_laser/sat \u2248 2,83\u00d710\u00b9\u00b2 W")

    pic(doc, os.path.join(DIAG, "diagram_lagrange.png"), w=7.5,
        cap="Fig. 2 \u2014 Punctele Lagrange L4 \u015fi L5 \u00een sistemul Soare\u2013P\u0103m\u00e2nt")

    # ── Cap. 3 ────────────────────────────────────────────────────────────────
    hd(doc, "3. Originalitatea \u015fi Inova\u021bia Proiectului", lvl=1, sz=12, sb=6, sa=2)
    tx(doc,
       "P.O.L.A.R.I.S. integreaz\u0103 patru domenii avansate \u00eentr-un lan\u021b coerent: "
       "ob\u021binerea energiei solare printr-un roi Dyson de oglinzi (\u224830 oglinzi cu \u00d81 km "
       "la 7 mln km de Soare), conversia \u00een laser de mare putere (\u03b7 = 43%) de 4 satelit\u021bi "
       "colectori, transferul energetic prin noduri la L4/L5, \u015fi transportul terestru prin "
       "cabluri YBCO sus\u021binute de baloane la 30 km altitudine.")
    tx(doc,
       "Selec\u021bia YBCO fa\u021b\u0103 de NbTi clasic constituie inova\u021bia tehnic\u0103 cheie. "
       "Discu\u021biile cu fizicieni de la CERN au confirmat c\u0103 YBCO opereaz\u0103 la 77 K "
       "cu azot lichid (nu heliu la 4 K ca NbTi din magne\u021bii LHC), reducând drastic "
       "complexitatea criogenic\u0103 \u015fi masa sistemului. Pe parcursul proiectului se a\u015fteapt\u0103 "
       "transmiterea a p\u00e2n\u0103 la 800 GW per cablu.")
    tx(doc,
       "Tranzi\u021bia Kardashev: prin arhitectura propus\u0103, P.O.L.A.R.I.S. adreseaz\u0103 "
       "avansul civiliza\u021biei umane pe scala Kardashev de la Tipul 0,73 (prezent) "
       "spre Tipul 1,0 \u2014 prima civiliza\u021bie care st\u0103p\u00e2ne\u015fte integral energia "
       "propriei planete. Aceasta nu este doar un calcul de inginerie, ci o viziune "
       "civiliza\u021bional\u0103 coerent\u0103.")

    pic(doc, os.path.join(IMG, "Image J - YBCO Superconductor Crystal.png"), w=6.0,
        cap="Fig. 3 \u2014 Cristal supraconductor YBCO \u2014 inova\u021bia de material a proiectului P.O.L.A.R.I.S.")

    # ── Cap. 4 ────────────────────────────────────────────────────────────────
    hd(doc, "4. Obiectivele Cercet\u0103rii", lvl=1, sz=12, sb=6, sa=2)
    hd(doc, "4.1 Obiectiv General", lvl=2, sz=11, sb=3, sa=1)
    tx(doc,
       "Analiza teoretic\u0103 \u015fi demonstrarea coerensei unui SES capabil s\u0103 livreze "
       "1,6 TW continuu c\u0103tre P\u0103m\u00e2nt, \u00een limitele legilor fizicii cunoscute.")
    hd(doc, "4.2 Obiective Specifice", lvl=2, sz=11, sb=3, sa=1)
    for obj in [
        "(1) Determinarea puterii solare captate \u015fi a puterii laser disponibile \u00een nodurile L4/L5.",
        "(2) Calcularea puterii maxime transmisibile printr-un cablu YBCO proiectat pentru ~800 GW.",
        "(3) Verificarea flotabilit\u0103\u021bii la 30 km altitudine (\u224811 t cu baloane de heliu).",
        "(4) Evaluarea impactului energetic global (procentaj din consumul mondial 2023).",
    ]:
        blt(doc, obj, sz=10)

    # ── Cap. 5 ────────────────────────────────────────────────────────────────
    hd(doc, "5. Ipoteze \u015fi Predic\u021bii", lvl=1, sz=12, sb=6, sa=2)
    tx(doc,
       "Ipoteza principal\u0103: Un SES format din roi Dyson + satelit\u021bi colectori + noduri "
       "Lagrange + sta\u021bii polare YBCO poate transmite 2 \u00d7 800 GW (total 1,6 TW) spre P\u0103m\u00e2nt, "
       "cu coeren\u021b\u0103 fizic\u0103 verificabil\u0103 \u015fi materiale aflate la gran\u021ba cercet\u0103rii actuale.")
    tx(doc,
       "Predic\u021bie: Puterea disponibil\u0103 la nivelul L4/L5 va dep\u0103\u015fi necesarul celor "
       "dou\u0103 centrale polare (\u224847% utilizat), justific\u00e2nd logic extensia Star Power Grid "
       "din surplusul energetic r\u0103mas.")

    # ── Cap. 6 ────────────────────────────────────────────────────────────────
    hd(doc, "6. Materiale \u015fi Metode", lvl=1, sz=12, sb=6, sa=2)
    hd(doc, "6.1 Materiale", lvl=2, sz=11, sb=3, sa=1)
    tx(doc, "Resurse teoretice:", bold=True, sz=10, sa=1)
    blt(doc, "Studii despre Space-Based Solar Power (SBSP/SPS) \u2014 NASA, JAXA, ESA", sz=10)
    blt(doc, "Documenta\u021bie CERN privind supraconductori HTS \u015fi criogenie", sz=10)
    blt(doc, "Studii HTS-YBCO pentru aplica\u021bii electrice (Larbalestier et al., 2001)", sz=10)
    tx(doc, "Software:", bold=True, sz=10, sa=1)
    blt(doc, "Excel / Sheets: calcule numerice ale fluxului energetic", sz=10)
    blt(doc, "Word, PowerPoint: redactare \u015fi structurare proiect", sz=10)
    blt(doc, "Kling 3.0: generarea celor 15 anima\u021bii video AI integrate \u00een prezentare", sz=10)
    blt(doc, "Gemini Imagen: generarea vizualiz\u0103rilor arhitecturale ale componentelor", sz=10)
    tx(doc, "Prezentare vizual\u0103 digital\u0103 (exponat multimedia):", bold=True, sz=10, sa=1)
    blt(doc,
        "Prezentare PowerPoint P.O.L.A.R.I.S. \u2014 20 slide-uri cu vizualiz\u0103ri "
        "cinematografice AI (Kling 3.0), incluz\u00e2nd 15 anima\u021bii video integrate", sz=10)
    blt(doc,
        "Anima\u021bii 3D ale componentelor: Dyson swarm, satelit\u021bi colectori, "
        "noduri L4/L5, sta\u021bii polare, re\u021bea Star Power Grid", sz=10)
    blt(doc,
        "Material audiovizual 45 secunde pentru afi\u015faj TV 50\u2033 la stand", sz=10)
    blt(doc,
        "Materiale imprimate: roll-up 80\u00d7200 cm, display carton 120\u00d760 cm", sz=10)

    hd(doc, "6.2 Metodologie", lvl=2, sz=11, sb=3, sa=1)
    for step in [
        "(1) Documentare \u015fi definire parametri fizici (constante, geometrie, randamente).",
        "(2) Modelare matematic\u0103 a fluxului energetic de la oglinzi p\u00e2n\u0103 la cablurile YBCO.",
        "(3) Dimensionare cablu supraconductor YBCO: P = U \u00d7 I; j = I / A \u2264 j_c.",
        "(4) Modelare flotabilitate baloane: F_A = \u03c1_aer \u00d7 V \u00d7 g la 30 km altitudine.",
        "(5) Modelare vizual\u0103 3D \u015fi corelare cu extensia Star Power Grid.",
    ]:
        blt(doc, step, sz=10)

    # ── Cap. 7 ────────────────────────────────────────────────────────────────
    hd(doc, "7. Arhitectura Sistemului Energetic", lvl=1, sz=12, sb=6, sa=2)
    tx(doc,
       "Nivel Solar \u2014 Roiul Dyson de Oglinzi: 30 de oglinzi cu diametru 1 km, "
       "amplasate \u00een centur\u0103 circumsolar\u0103 la ~7 milioane km de Soare. Reflectivitate 90%. "
       "Putere interceptat\u0103 total\u0103: \u22481,46 \u00d7 10\u00b9\u00b3 W.")
    tx(doc,
       "Colectori Laser: 4 satelit\u021bi colectori la ~10 milioane km. Fiecare capteaz\u0103 "
       "P_lum/sat \u2248 3,30 \u00d7 10\u00b9\u00b2 W \u015fi o converte\u015fte \u00een laser cu \u03b7 = 43%, "
       "gener\u00e2nd P_laser/sat \u2248 1,42 \u00d7 10\u00b9\u00b2 W.")
    tx(doc,
       "Noduri Lagrange L4 \u015fi L5: c\u00e2te 2 satelit\u021bi colectori direc\u021bioneaz\u0103 laserul "
       "spre fiecare nod. Putere cumulat\u0103 per nod: \u22482,83 \u00d7 10\u00b9\u00b2 W. Surplusul "
       "fa\u021b\u0103 de necesarul centralei polare (~53% din disponibil) alimenteaz\u0103 extensia "
       "Star Power Grid.")
    tx(doc,
       "Sta\u021bii orbitale polare: centrale electrice sus\u021binute de baloane cu heliu la ~30 km "
       "altitudine. Mas\u0103 total\u0103 per ansamblu: \u224811 t (central\u0103 \u22644 t + cablu YBCO \u22487 t). "
       "Putere electric\u0103 produs\u0103: 800 GW per sta\u021bie, transmis\u0103 la sol prin cabluri YBCO.")

    pic(doc, os.path.join(IMG, "00 System Diagram.png"), w=14.0,
        cap="Fig. 4 \u2014 Diagrama complet\u0103 a sistemului P.O.L.A.R.I.S. "
            "(roi Dyson \u2192 colectori \u2192 L4/L5 \u2192 sta\u021bii polare \u2192 YBCO \u2192 P\u0103m\u00e2nt)")

    # ── Cap. 8 ────────────────────────────────────────────────────────────────
    hd(doc, "8. STAR POWER GRID \u2014 Sintez\u0103", lvl=1, sz=12, sb=6, sa=2)
    tx(doc,
       "Star Power Grid este extensia terestr\u0103 a sistemului P.O.L.A.R.I.S., proiectat\u0103 "
       "pe termen mediu de 30\u201350 de ani. Energia recep\u021bionat\u0103 la sta\u021biile orbitale polare "
       "(Polul Nord \u015fi Polul Sud) este transmis\u0103 prin cablurile YBCO p\u00e2n\u0103 la capetele "
       "de re\u021bea terestre \u2014 cele dou\u0103 mari noduri de distribu\u021bie ale sistemului.")
    tx(doc,
       "De la cei doi poli, distribu\u021bia se realizeaz\u0103 radiant spre continente "
       "printr-o re\u021bea de linii de \u00eenalt\u0103 tensiune supraconductoare, similar\u0103 re\u021belelor "
       "HVDC existente, dar la scar\u0103 planetar\u0103. Fiecare pol devine un hub energetic "
       "de 800 GW care alimenteaz\u0103 o re\u021bea de transmisie terestr\u0103 cu acoperire global\u0103.")
    tx(doc,
       "Arhitectura bipolar\u0103: distribu\u021bie simetric\u0103 de la Polul Nord (spre Europa, Asia, "
       "America de Nord) \u015fi de la Polul Sud (spre America de Sud, Africa, Oceania). "
       "Aceast\u0103 configura\u021bie minimizeaz\u0103 lungimea total\u0103 a re\u021belei \u015fi asigur\u0103 "
       "acces la energie pentru regiunile izolate geografic.")

    pic(doc, os.path.join(IMG, "Image I - Star Power Grid GEO.png"), w=10.0,
        cap="Fig. 5 \u2014 Star Power Grid: distribu\u021bie radiant\u0103 de la poli spre continente "
            "(cabluri terestre globale de \u00eenalt\u0103 tensiune)")

    # ── Cap. 9 ────────────────────────────────────────────────────────────────
    hd(doc, "9. Rezultate \u2014 Fluxul Energetic Complet", lvl=1, sz=12, sb=6, sa=2)
    hd(doc, "9.1 Date brute (input)", lvl=2, sz=11, sb=3, sa=1)
    tx(doc,
       "Constante fizice: I_Earth = 1.361 W/m\u00b2, d_SE = 1,496\u00d710\u00b9\u00b9 m. "
       "Geometrie: 30 oglinzi, \u00d81 km, la 7 mln km de Soare. 4 satelit\u021bi colectori. "
       "Randamente: reflectivitate 90%, \u03b7_laser 43%, \u03b7_electric 60%.")

    hd(doc, "9.2\u20139.7 Fluxul energetic \u2014 sinteza rezultatelor", lvl=2, sz=11, sb=3, sa=2)
    rows9 = [
        ("Putere inciden\u021b\u0103 pe oglinzi",       "\u22481,46 \u00d7 10\u00b9\u00b3 W"),
        ("Putere reflectat\u0103 (90%)",                "\u22481,32 \u00d7 10\u00b9\u00b3 W"),
        ("Lumin\u0103 per colector (1/4)",              "\u22483,30 \u00d7 10\u00b9\u00b2 W"),
        ("Laser per colector (43%)",                    "\u22481,42 \u00d7 10\u00b9\u00b2 W"),
        ("Laser total (4 colectori)",                   "\u22485,67 \u00d7 10\u00b9\u00b2 W"),
        ("Putere \u00een L4 / L5 (fiecare nod)",        "\u22482,83 \u00d7 10\u00b9\u00b2 W"),
        ("Laser necesar per central\u0103 polar\u0103",  "\u22481,33 \u00d7 10\u00b9\u00b2 W (pt. 800 GW)"),
        ("Procentaj L4/L5 utilizat",                    "\u224847%"),
        ("Putere electric\u0103 total\u0103 livrat\u0103","1,6 TW (2 \u00d7 800 GW)"),
        ("Ponderea \u00een consumul electric global",   "\u224848% (fa\u021b\u0103 de 3,36 TW \u00een 2023)"),
    ]
    tbl9 = doc.add_table(rows=len(rows9)+1, cols=2)
    tbl9.style = 'Table Grid'
    # header
    h0 = tbl9.rows[0].cells[0]; h1 = tbl9.rows[0].cells[1]
    set_cw(h0, 9.0); set_cw(h1, 6.5)
    cell_tx(h0, "M\u0103rime / Etap\u0103 calcul", bold=True, sz=10)
    cell_tx(h1, "Valoare", bold=True, sz=10)
    for i, (label, val) in enumerate(rows9):
        r = tbl9.rows[i+1]
        set_cw(r.cells[0], 9.0); set_cw(r.cells[1], 6.5)
        cell_tx(r.cells[0], label, sz=10)
        cell_tx(r.cells[1], val, sz=10)
    doc.add_paragraph()

    pic(doc, os.path.join(DIAG, "diagram_energy_bar.png"), w=10.0,
        cap="Fig. 6 \u2014 Flux energetic complet: de la Soare la P\u0103m\u00e2nt")

    # ── Cap. 10 ───────────────────────────────────────────────────────────────
    hd(doc, "10. Estimare de Costuri", lvl=1, sz=12, sb=6, sa=2)
    tx(doc,
       "Aceast\u0103 sec\u021biune ofer\u0103 ordine de m\u0103rime, nu pre\u021buri exacte. "
       "Infrastructura solar\u0103 + satelit\u021bi + noduri L4/L5 + centrale: la 20 EUR/W "
       "pentru 1,6 TW \u2192 \u224832.000 miliarde EUR la pre\u021burile actuale. "
       "Cabluri YBCO (2\u00d730 km): 2\u201310 miliarde EUR; baloane + sus\u021binere: 1\u20134 miliarde EUR.")
    tx(doc,
       "ROI la scar\u0103 complet\u0103: energie valorificat\u0103 \u2248500\u2013600 miliarde EUR/an "
       "(la 40 USD/MWh). Economicitate realist\u0103 pe orizonturi lungi (50+ ani), cu "
       "ieftinire semnificativ\u0103 a tehnologiilor spa\u021biale \u015fi HTS.")

    # ── Cap. 11 ───────────────────────────────────────────────────────────────
    hd(doc, "11. Analiza Rezultatelor", lvl=1, sz=12, sb=6, sa=2)
    hd(doc, "11.1 Fluxul energetic de la Soare la L4/L5", lvl=2, sz=11, sb=3, sa=1)
    tx(doc,
       "Calculele sunt coerente cu legile fizicii. Chiar cu pierderi semnificative "
       "(\u03b7_laser = 43%, \u03b7_electric = 60%), puterea disponibil\u0103 \u00een L4/L5 "
       "dep\u0103\u015fe\u015fte cu mult necesarul celor dou\u0103 centrale de 800 GW, "
       "l\u0103s\u00e2nd o marj\u0103 consistent\u0103 pentru Star Power Grid.")
    hd(doc, "11.2 Limitarea la 2 \u00d7 800 GW prin YBCO", lvl=2, sz=11, sb=3, sa=1)
    tx(doc,
       "Cablurile dimensionate pentru 800 GW/cablu (j \u2264 j_c) cu mas\u0103 \u22487 t pe 30 km "
       "reprezint\u0103 o constr\u00e2ngere material\u0103 \u015fi inginereasc\u0103 realist\u0103, nu o limit\u0103 "
       "energetic\u0103. Limitarea este o alegere con\u015ftient\u0103 pentru a ancora proiectul "
       "\u00een parametri verificabili.")
    hd(doc, "11.3 Masa total\u0103 \u015fi flotabilitatea baloanelor", lvl=2, sz=11, sb=3, sa=1)
    tx(doc,
       "\u03c1_aer la 30 km \u2248 0,018 kg/m\u00b3 permite sus\u021binerea a 11 t cu volume de heliu "
       "de ordinul miilor de m\u00b3. Solu\u021bie la limita tehnicii actuale; necesit\u0103 materiale "
       "ultra-u\u015foare \u015fi sisteme avansate de control al stabilit\u0103\u021bii.")
    hd(doc, "11.4 Impactul energetic global", lvl=2, sz=11, sb=3, sa=1)
    tx(doc,
       "1,6 TW continuu = 48% din consumul electric global (2023: \u22483,36 TW mediu). "
       "Impact echivalent cu zeci de centrale nucleare sau o mare parte din produc\u021bia "
       "regenerabil\u0103 actual\u0103 \u2014 proiectul dep\u0103\u015fe\u015fte relevan\u021ba local\u0103 "
       "\u015fi are valoare global\u0103.")
    hd(doc, "11.5 Star Power Grid \u2014 analiz\u0103 conceptual\u0103", lvl=2, sz=11, sb=3, sa=1)
    tx(doc,
       "Surplusul energetic din L4/L5 (~53% neutilizat de centrale) poate alimenta "
       "noduri terestre suplimentare ale re\u021belei globale. Limitat mai mult de factori "
       "tehnologici (materiale, control) dec\u00e2t energetici.")
    hd(doc, "11.6 Tranzi\u021bia pe Scala Kardashev", lvl=2, sz=11, sb=3, sa=1)
    tx(doc,
       "Consumul uman actual (~18 TW total, ~3,36 TW electricitate) plaseaz\u0103 "
       "civiliza\u021bia la Tipul 0,73 pe scala Kardashev. P.O.L.A.R.I.S. contribuie cu "
       "1,6 TW electricitate spre Tipul 1,0 (controlul integral al energiei planetare "
       "\u2248 10\u00b9\u2077 W). Prin extensii progresive ale arhitecturii Dyson, sistemul "
       "poate scala spre aceast\u0103 limit\u0103, marc\u00e2nd o tranzi\u021bie civiliza\u021bional\u0103.")

    # Imagini Kardashev: F1 si F2 una dupa alta (compact)
    pic(doc, os.path.join(IMG, "Image F1 - Kardashev Earth Today.png"), w=7.0,
        cap="Fig. 7a \u2014 Tipul 0,73 Kardashev: P\u0103m\u00e2ntul \u00een prezent")
    pic(doc, os.path.join(IMG, "Image F2 - Kardashev Earth POLARIS (1.0).png"), w=7.0,
        cap="Fig. 7b \u2014 Tipul 1,0 Kardashev: P\u0103m\u00e2ntul cu P.O.L.A.R.I.S.")

    # ── Cap. 12 ───────────────────────────────────────────────────────────────
    hd(doc, "12. Discu\u021bii, Limit\u0103ri \u015fi Concluzii", lvl=1, sz=12, sb=6, sa=2)
    tx(doc,
       "Proiectul P.O.L.A.R.I.S. demonstreaz\u0103 c\u0103 un SES bazat pe roi Dyson, satelit\u021bi "
       "laser, noduri Lagrange \u015fi supraconductori YBCO este teoretic coerent, "
       "respect\u00e2nd legile fizicii cu materiale \u015fi tehnologii aflate la gran\u021ba cercet\u0103rii.")
    tx(doc,
       "Rezultatele principale: (1) Intensitate solar\u0103 \u22486,22\u00d710\u2075 W/m\u00b2 la 7 mln km. "
       "(2) Putere disponibil\u0103 \u00een L4/L5: \u22482,83 TW per nod. "
       "(3) 1,6 TW livra\u021bi continuu = 48% din electricitatea global\u0103. "
       "(4) YBCO validat pentru transport de 800 GW la 77 K cu mas\u0103 minim\u0103.")
    tx(doc,
       "Limitele proiectului sunt identificate explicit: distan\u021bele solare extreme, "
       "cabluri YBCO f\u0103r\u0103 precedent, baloane stratosferice mega-scale, "
       "provoc\u0103ri de siguran\u021b\u0103 \u015fi reglementare interna\u021bional\u0103. "
       "Nu propunem o solu\u021bie imediat implementabil\u0103, ci un model teoretic riguros "
       "\u015fi o viziune coerent\u0103 pentru infrastructura energetic\u0103 a viitorului.")

    # ── Bibliografie ──────────────────────────────────────────────────────────
    hd(doc, "Bibliografie", lvl=1, sz=12, sb=6, sa=2)
    refs = [
        "[1] Glaser, P.E. (1968). Power from the Sun: Its Future. Science, 162(3856), 857\u2013861. doi:10.1126/science.162.3856.857",
        "[2] IEA (2024). Electricity 2024: Analysis and forecast to 2026. https://www.iea.org/reports/electricity-2024",
        "[3] JAXA (n.d.). Space Solar Power Systems (SSPS). https://www.kenkai.jaxa.jp/eng/research/ssps/ssps-index.html",
        "[4] Mankins, J.C. (2011). SPS-ALPHA. NASA NIAC Phase I Report. https://www.researchgate.net/publication/268573928",
        "[5] NASA (n.d.). Space-Based Solar Power. https://www.nasa.gov/organizations/otps/space-based-solar-power-report/",
        "[6] NASA GSFC (n.d.). Total Solar Irradiance. https://solarscience.msfc.nasa.gov/",
        "[7] NASA (n.d.). What are Lagrange Points? https://solarsystem.nasa.gov/resources/754/",
        "[8] NASA (2018). Parker Solar Probe. https://www.nasa.gov/content/goddard/parker-solar-probe",
        "[9] Brown, W.C. (1984). History of Power Transmission by Radio Waves. IEEE TMT, 32(9), 1230\u20131242. doi:10.1109/TMTT.1984.1132833",
        "[10] Shinohara, N. (2014). Wireless Power Transfer via Radiowaves. ISTE/Wiley. ISBN: 978-1-848-21518-1",
        "[11] CERN (2021). High-temperature superconductors for future accelerators. https://home.cern/science/engineering/superconductivity",
        "[12] Larbalestier, D. et al. (2001). High-Tc superconducting materials for electric power. Nature, 414, 368\u2013377. doi:10.1038/35104654",
    ]
    for ref in refs:
        pp = doc.add_paragraph(style='List Number')
        pp.paragraph_format.space_after = Pt(2)
        rr = pp.add_run(ref[4:])   # skip "[N] " prefix — list style adds number
        rr.font.size = Pt(9.5)
        rr.font.name = 'Calibri'

    # ── Ultima pagina: Autori + Coordonator + Colaboratori ───────────────────
    doc.add_page_break()
    hd(doc, "Echip\u0103 \u015fi Colaboratori", lvl=1, sz=12, sb=4, sa=6)

    team_rows = [
        ("Ioan Cristian CHELARU",
         "Clasa a VIII-a | \u26a0 Unitate \u015fcolar\u0103: de completat\n"
         "Rol: Modelare vizual\u0103, scala Kardashev, re\u021beaua Star Power Grid, prezentare"),
        ("Albert David OLARIU",
         "Clasa a VIII-a | \u26a0 Unitate \u015fcolar\u0103: de completat\n"
         "Rol: Arhitectur\u0103 sistem, calcule energetice, supraconductori YBCO, dinamic\u0103 orbital\u0103"),
        ("Prof. Francisc Dionisie Aaron\n(Mentor / Coordonator)",
         "ICHB \u2014 Liceul Teoretic Interna\u021bional de Informatic\u0103\n"
         "\u26a0 Nr. telefon: de completat"),
        ("Daniel-Justinian ZELENSCHI\n(Colaborator)",
         "Ing. chimic (UCL + Melbourne) \u2022 PhD(c) AI \u2022 implementare AI la scar\u0103 industrial\u0103 (Raiffeisen AI) \u2022 NASA Ames Research Center \u2014 Space Settlement Contest \u2014 Premiul II\n"
         "Consultan\u021b\u0103: algoritmi AI pentru controlul autonom al swarm-ului orbital \u015fi optimizarea Star Power Grid"),
        ("Carlo Emilio MONTANARI\n(Colaborator)",
         "Fizician aplicat \u2022 Doctorand Dinamica Fasciculelor, CERN / Univ. Bologna\n"
         "Consultan\u021b\u0103: dinamica fasciculelor de \u00eenalt\u0103 precizie (LHC \u2192 POLARIS)"),
        ("Prof. Univ. Dr. Valentin BARNA\n(Colaborator)",
         "Facultatea de Fizic\u0103, Universitatea din Bucure\u015fti\n"
         "Specializare: Optic\u0103 \u015fi Fotonic\u0103, Fizica LASER-ilor\n"
         "Consultan\u021b\u0103: principii de fotonic\u0103 \u015fi conversie radia\u021bie solar\u0103 \u2192 LASER "
         "aplicabile lan\u021bului energetic"),
    ]

    tbl_t = doc.add_table(rows=len(team_rows), cols=2)
    tbl_t.style = 'Table Grid'
    for i, (name, detail) in enumerate(team_rows):
        r = tbl_t.rows[i]
        c0 = r.cells[0]; c1 = r.cells[1]
        set_cw(c0, 5.5); set_cw(c1, 9.5)

        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(3)
        r0 = p0.add_run(name)
        r0.bold = True; r0.font.size = Pt(10); r0.font.name = 'Calibri'

        lines = detail.split('\n')
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(1)
        rr1 = p1.add_run(lines[0])
        rr1.font.size = Pt(10); rr1.font.name = 'Calibri'
        for line in lines[1:]:
            pp = c1.add_paragraph()
            pp.paragraph_format.space_after = Pt(1)
            rr = pp.add_run(line)
            rr.font.size = Pt(9.5); rr.font.name = 'Calibri'
            rr.italic = '\u26a0' in line

    doc.add_paragraph()
    pic(doc, os.path.join(ASS, "team-hero.jpeg"), w=8.0,
        cap="Echipa POLARIS Bears \u2014 Ioan Cristian CHELARU \u015fi Albert David OLARIU")

    path = os.path.join(OUT, "05_Document_Principal.docx")
    doc.save(path)
    print("  [OK] 05_Document_Principal.docx")

# ─── MAIN ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import io, sys
    if hasattr(sys.stdout, 'buffer'):
        sys.stdout = io.TextIOWrapper(
            sys.stdout.buffer, encoding='utf-8', errors='replace')

    os.makedirs(OUT, exist_ok=True)
    print("Generare documente ONCS 2026 / P.O.L.A.R.I.S. ...")
    doc1_coperta()
    doc2_cuprins()
    doc3_identificare()
    doc4_rezumat()
    doc5_principal()
    print("\nGata! 5 fisiere salvate in:", OUT)

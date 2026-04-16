import os
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not found. Please install it with 'pip install python-docx'")
    exit(1)

def create_bibliography():
    doc = Document()

    # Title
    title = doc.add_heading('POLARIS INDEX', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Scientific References', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # The text to insert
    references = [
        ("01", "Glaser, P.E., 1968. Power from the Sun: Its Future. Science, 162(3856), pp.857–861. doi.org/10.1126/science.162.3856.857", "FOUNDATIONAL — SBSP CONCEPT"),
        ("02", "International Energy Agency (IEA), 2024. Electricity 2024: Analysis and forecast to 2026. iea.org/reports/electricity-2024", "GLOBAL ENERGY DATA — 3.36 TW REFERENCE"),
        ("03", "Japan Aerospace Exploration Agency (JAXA), n.d. Space Solar Power Systems (SSPS). kenkai.jaxa.jp/eng/research/ssps", "MICROWAVE ENERGY TRANSMISSION — KW SCALE DEMO"),
        ("04", "Mankins, J.C., 2011. SPS-ALPHA: The First Practical Solar Power Satellite via Arbitrarily Large Phased Array. NASA NIAC Phase I Report. ntrs.nasa.gov — NASA Technical Reports Server", "SPS ARCHITECTURE — PHASED ARRAY"),
        ("05", "NASA, n.d. Space-Based Solar Power. NASA Innovative Advanced Concepts (NIAC). nasa.gov/niac", "NASA SBSP PROGRAM"),
        ("06", "NASA Marshall Space Flight Center, n.d. Solar Physics — Total Solar Irradiance. solarscience.msfc.nasa.gov", "SOLAR CONSTANT — 1361 W/m² REFERENCE"),
        ("07", "NASA, n.d. What are Lagrange Points? NASA Science — Solar System Exploration. science.nasa.gov/solar-system/lagrange-points", "L4/L5 LAGRANGE NODES"),
        ("08", "NASA, 2018. Parker Solar Probe Mission Overview. science.nasa.gov/mission/parker-solar-probe", "PROOF — PROXIMITY TO SUN IS FEASIBLE (6.1M KM)"),
        ("09", "Brown, W.C., 1984. The History of Power Transmission by Radio Waves. IEEE Transactions on Microwave Theory and Techniques, 32(9), pp.1230–1242.", "WIRELESS POWER TRANSMISSION — HISTORICAL BASIS"),
        ("10", "Shinohara, N., 2014. Wireless Power Transfer via Radiowaves. London: ISTE Ltd.", "STAR POWER GRID — MICROWAVE TRANSMISSION BASIS"),
        ("11", "CERN, 2021. High-temperature superconductors for future accelerators. home.cern/science/engineering/superconductivity", "YBCO vs NbTi — CERN CONSULTATION BASIS"),
        ("12", "Larbalestier, D., Gurevich, A., Feldmann, D.M. and Polyanskii, A., 2001. High-T c superconducting materials for electric power applications. Nature, 414(6861), pp.368–377. doi.org/10.1038/35104654", "HTS/YBCO MATERIALS SCIENCE")
    ]

    for index, text, label in references:
        p = doc.add_paragraph()
        run_index = p.add_run(f"[{index}] ")
        run_index.bold = True
        run_index.font.size = Pt(12)
        
        run_text = p.add_run(text)
        run_text.font.size = Pt(11)
        
        p_label = doc.add_paragraph()
        run_label = p_label.add_run(f"   ↳ {label}")
        run_label.italic = True
        run_label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run_label.font.size = Pt(9)
        
        doc.add_paragraph() # Spacer

    # Footer section
    doc.add_page_break()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_collab = footer_p.add_run("Collaboratori Polaris Bears // ONCS 2026\nChelaru Ioan Cristian & Olariu David Albert")
    run_collab.bold = True
    
    footer_p2 = doc.add_paragraph()
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_final = footer_p2.add_run("POLARIS COMMAND © 2026\nALL SYSTEMS NOMINAL.")
    run_final.font.size = Pt(10)
    run_final.font.color.rgb = RGBColor(0x00, 0x88, 0x00)

    file_path = "POLARIS_Bibliography.docx"
    doc.save(file_path)
    print(f"File saved successfully at: {os.path.abspath(file_path)}")

if __name__ == "__main__":
    create_bibliography()

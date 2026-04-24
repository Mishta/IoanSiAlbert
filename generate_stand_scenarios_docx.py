"""Generates the Word document with stand evaluation scenarios for ONCS 2026.

Format: A4 portrait, Celestial Command palette.
Output: IoanSiAlbert/docs/Scenarii_Evaluare_Stand_ONCS_2026.docx
"""
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Antigravity\__SKILLS\IoanSiAlbert\docs\Scenarii_Evaluare_Stand_ONCS_2026.docx"
SLOGAN = "Inovăm prezentul. Alimentăm viitorul!"

# Celestial Command palette
NAVY = RGBColor(0x0F, 0x13, 0x1F)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
GOLD_BRIGHT = RGBColor(0xFF, 0xBA, 0x20)
CYAN = RGBColor(0x00, 0x88, 0xAA)
GREY_LIGHT = RGBColor(0x55, 0x55, 0x55)
GREY_VLIGHT = RGBColor(0x88, 0x88, 0x88)

# ============== HELPERS ==============

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def para(doc, text, *, bold=False, italic=False, size=11, color=None,
         align=None, space_after=6, space_before=0, font='Calibri',
         first_line_indent=None, left_indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p

def heading(doc, text, *, level=1, color=None):
    if color is None:
        color = GOLD if level == 1 else NAVY
    size = {1: 18, 2: 14, 3: 12}.get(level, 11)
    space_before = {1: 6, 2: 12, 3: 8}.get(level, 4)
    return para(doc, text, bold=True, size=size, color=color,
                space_before=space_before, space_after=6)

def bullet(doc, text, *, size=10.5, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    return p

def quote_block(doc, text, *, color=None, italic=True, size=11):
    """Indented quotation-style paragraph with left border."""
    if color is None:
        color = NAVY
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.italic = italic
    r.font.color.rgb = color
    # add left border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), 'C9A84C')
    pBdr.append(left)
    pPr.append(pBdr)
    return p

def mono_block(doc, text, *, size=9):
    """Monospace block for ASCII diagrams."""
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(size)
        r.font.color.rgb = NAVY

def make_table(doc, data, widths_cm, *, header_row=True, size=10):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.autofit = False
    widths = [Cm(w) for w in widths_cm]
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.width = widths[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            # cell may have multi-line content
            first = True
            for line in str(val).split('\n'):
                p = cell.paragraphs[0] if first else cell.add_paragraph()
                first = False
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(line)
                run.font.name = 'Calibri'
                run.font.size = Pt(size)
                if header_row and i == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if header_row and i == 0:
                set_cell_shading(cell, '0F131F')
            else:
                set_cell_shading(cell, 'F5F5F5' if i % 2 else 'FFFFFF')
    return table

def page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

# ============== DOCUMENT ==============

doc = Document()

# A4 explicit — required per project standard
for section in doc.sections:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ============== COVER ==============

para(doc, "", size=11, space_after=24)
para(doc, "P.O.L.A.R.I.S.", bold=True, size=32, color=GOLD,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "Platformă Orbitală Laser pentru Alimentare și Recepție Inovativă Solară",
     italic=True, size=11, color=GREY_LIGHT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

para(doc, "Scenarii de Evaluare la Stand",
     bold=True, size=22, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "Ghid operațional pentru juriul ONCS 2026",
     italic=True, size=13, color=CYAN,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

para(doc, "Echipa POLARIS Bears",
     bold=True, size=13, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "Albert David OLARIU  ·  Ioan Cristian CHELARU",
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "Clasa a VIII-a  ·  Secțiunea C — Tehnologia Informației",
     italic=True, size=10, color=GREY_LIGHT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "Mentor: Prof. Dr. Francisc Dionisie Aaron",
     italic=True, size=10, color=GREY_LIGHT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)

# SLOGAN — cover
para(doc, f"„{SLOGAN}\"",
     bold=True, italic=True, size=16, color=GOLD_BRIGHT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "— sloganul echipei —",
     italic=True, size=9, color=GREY_VLIGHT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

page_break(doc)

# ============== SECȚIUNEA 1 — CONTEXT ==============

heading(doc, "1. Contextul evaluării la stand — ONCS 2026", level=1)
para(doc,
     "Conform Regulamentului ONCS (Art. 16 — Evaluarea proiectelor), juriul studiază mai întâi "
     "Fișa proiectului și displayul/exponatul (~15 minute), apoi participă la prezentarea "
     "echipei (10 minute) urmată de sesiunea de întrebări. La evaluarea la stand, juriul se "
     "deplasează la exponatul vostru — nu invers. Mentorii, însoțitorii și părinții nu sunt "
     "prezenți. Elevii vor fi legitimați.",
     size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)

para(doc,
     "Există două situații reale pe care trebuie să le pregătiți separat:",
     size=11, space_after=6)

bullet(doc, "SCENARIUL A — Evaluare formală la stand (programată, 15–25 min, juriu complet)")
bullet(doc, "SCENARIUL B — Evaluare flash / juror solo (neanunțată, 3–8 min, informal)")

para(doc, "", space_after=4)
heading(doc, "Materiale disponibile la stand", level=2)

materiale = [
    ("Material", "Detalii", "Rol în evaluare"),
    ("TV cu video loop", "Animația P.O.L.A.R.I.S. cu audio, buclă continuă\n(CLIP_POLARIS_Sequence_audio.mp4)",
     "Backdrop vizual permanent + cârlig pentru vizitatori"),
    ("Display oficial ONCS", "Triptych 30 + 60 + 30 cm (vertical)\nDocumentația proiectului — regulament",
     "Referință principală pentru juriu; conține calcule și scheme"),
    ("Roll-up marketing", "85 × 200 cm, vertical\n(rollup_4_complet.pdf)",
     "Identitate vizuală brand; NU înlocuiește displayul oficial"),
    ("Laptop + demo DiagramZoom", "Animația interactivă de pe polaris-bears.ro",
     "Demonstrație live (90 sec max în Scenariul A)"),
    ("Flyere", "Variante multiple cu QR code către website",
     "Take-away pentru juror; NU împărțite proactiv"),
    ("Mingi antistress", "Cu logo-ul POLARIS Bears",
     "DOAR pentru vizitatori non-juriu; nu în timpul evaluării"),
    ("Website live", "polaris-bears.ro (GA4 activ) + QR pe postere",
     "Aprofundare după prezentare"),
]
make_table(doc, materiale, [3.8, 6.5, 6.5])
para(doc, "", space_after=10)

page_break(doc)

# ============== SECȚIUNEA 2 — PITCH-URI ==============

heading(doc, "2. Pitch-ul de deschidere — 3 variante calibrate", level=1)
para(doc,
     "Toate cele trei variante de mai jos sunt pregătite ad litteram. Alegeți varianta "
     "potrivită momentului — în funcție de cine este în fața voastră și cât timp aveți. "
     "Semnul // marchează o pauză scurtă (0,5–1 sec) pentru ritm și emfază. "
     "Privirea merge către juriu, niciodată pe foaie.",
     size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

# ---------- Varianta A ----------
heading(doc, "Varianta A — Inspirațional (90 secunde)  ★ RECOMANDATĂ", level=2)
para(doc, "Ritm rar, cu pauze marcate. Privire la comisie, nu la foaie.",
     italic=True, size=10, color=GREY_LIGHT, space_after=8)

quote_block(doc,
    "„Bună ziua. Suntem POLARIS Bears.\n\n"
    "Trăim pe o planetă care consumă 20 terawatt-ore pe oră — // și totuși 70% din "
    "această energie vine de la combustibili care ne distrug clima. Întrebarea pe care "
    "ne-am pus-o este simplă: // dacă ar exista o sursă de energie de 457 de ori mai "
    "intensă decât cea solară terestră, continuă 24 din 24, fără nori, fără nopți, // "
    "am avea curajul să o accesăm?\n\n"
    "Răspunsul nostru se numește P.O.L.A.R.I.S. // — un sistem teoretic format din 30 "
    "de oglinzi orbitale la 7,7 milioane de kilometri de Soare, // 4 sateliți care "
    "convertesc lumina în laser, // noduri de transfer la punctele Lagrange L4 și L5, "
    "// și două stații polare conectate la Pământ prin cabluri supraconductoare YBCO. "
    "// Livrare finală: 1,6 terawați de electricitate continuă.\n\n"
    "Ce înseamnă asta? // 48% din consumul electric actual al întregii planete. // Un "
    "salt de la Tipul 0,73 spre Tipul 1,0 pe scara Kardashev — // prima civilizație "
    "capabilă să-și stăpânească propria planetă.\n\n"
    "Nu vă promitem că vom construi acest sistem mâine. // Vă demonstrăm că legile "
    "fizicii îl permit. // Și asta — este începutul.\""
)
para(doc, "Durată estimată: 85–90 secunde la ritm natural.",
     italic=True, size=9, color=GREY_VLIGHT, space_after=14)

# ---------- Varianta B ----------
heading(doc, "Varianta B — Tehnic-Direct (60 secunde)", level=2)
para(doc, "Pentru o comisie formată majoritar din fizicieni / ingineri. Ritm mai rapid.",
     italic=True, size=10, color=GREY_LIGHT, space_after=8)

quote_block(doc,
    "„Doamnelor și domnilor, proiectul P.O.L.A.R.I.S. răspunde la o singură întrebare: "
    "pot legile fizicii cunoscute să susțină un sistem energetic spațial care să livreze "
    "1,6 TW continuu către Pământ?\n\n"
    "Răspunsul nostru, documentat pe 20 de pagini cu calcule validate: DA.\n\n"
    "Arhitectura are patru niveluri: un roi Dyson de 30 oglinzi la 7,696 milioane km de "
    "Soare, unde intensitatea radiației este de 5,14 × 10⁵ W/m². // Patru sateliți "
    "colectori convertesc lumina în laser cu randament proiectat de 43%. // Nodurile "
    "Lagrange L4 și L5 redirecționează fasciculele. // Două stații polare terestre, "
    "legate prin cabluri YBCO supraconductoare la 77 Kelvin, distribuie 800 GW fiecare.\n\n"
    "Calculele sunt consistente. // Materialele — YBCO, oglinzi multistrat, convertoare "
    "laser-electricitate la 60% — există astăzi la scară de laborator. // Extrapolarea "
    "la scara propusă este frontiera cercetării actuale.\n\n"
    "Nu este inginerie pentru 2030. Este un model teoretic pentru secolul XXII. // "
    "Suntem gata pentru întrebări.\""
)
para(doc, "Durată estimată: 58–62 secunde.",
     italic=True, size=9, color=GREY_VLIGHT, space_after=14)

# ---------- Varianta C ----------
heading(doc, "Varianta C — Narativ (75 secunde)  ★ ALTERNATIVĂ", level=2)
para(doc, "Creează legătură emoțională cu comisia. Pentru juriu mixt (nu doar fizicieni).",
     italic=True, size=10, color=GREY_LIGHT, space_after=8)

quote_block(doc,
    "„În 1959, Freeman Dyson a propus o idee nebună: o civilizație suficient de avansată "
    "își poate înconjura propria stea pentru a capta toată energia acesteia. // A fost "
    "ridiculizat.\n\n"
    "În 1964, Nikolai Kardashev a pus pe hârtie o scală prin care putem măsura avansul "
    "unei civilizații după puterea pe care o controlează. // Pământul este astăzi la "
    "Tipul 0,73.\n\n"
    "În 2026, noi — Albert David Olariu și Ioan Cristian Chelaru, elevi în clasa a "
    "VIII-a la ICHB — am vrut să știm: cum ar arăta primul pas real către Tipul 1?\n\n"
    "Am petrecut luni întregi cu ecuații, consultări la CERN și la Universitatea "
    "București. // Am construit un model în care 30 de oglinzi orbitează Soarele la "
    "7,7 milioane km, convertesc lumina în laser, o trimit în Lagrange L4 și L5, și o "
    "aduc pe Pământ prin cabluri supraconductoare.\n\n"
    "Rezultatul: P.O.L.A.R.I.S. — un model teoretic care demonstrează că 1,6 terawați "
    "— 48% din electricitatea globală — nu este o fantezie, ci o consecință a legilor "
    "pe care deja le cunoaștem.\n\n"
    "Astăzi vă prezentăm nu o soluție pentru 2030. Ci o hartă pentru 2130.\""
)
para(doc, "Durată estimată: 72–78 secunde.",
     italic=True, size=9, color=GREY_VLIGHT, space_after=14)

# ---------- Tabel decizie ----------
heading(doc, "Când alegi ce variantă", level=2)
decizie = [
    ("Situație", "Varianta recomandată", "De ce"),
    ("Juriu complet la stand (Scenariul A), timp de 10 min", "Varianta A (Inspirațional)",
     "Context emoțional + cifră-stat imediat; merită cele 90 secunde"),
    ("Juriu majoritar cu profil literar / pedagogic / management", "Varianta C (Narativ)",
     "Arc istoric Dyson → Kardashev → 2026; captează audiența non-tehnică"),
    ("Juriu majoritar fizicieni / ingineri la stand formal", "Varianta A sau B",
     "A dacă ai timp; B dacă juriul e vizibil grăbit"),
    ("Flash — juror solo, 3–8 min (Scenariul B)", "Varianta B sau template „contract + pitch\"",
     "Vezi Secțiunea 4 — decideți pe loc, în funcție de semnale"),
]
make_table(doc, decizie, [5.2, 4.8, 6.8])
para(doc, "", space_after=6)

page_break(doc)

# ============== SECȚIUNEA 3 — SCENARIUL A ==============

heading(doc, "3. SCENARIUL A — Evaluare Formală la Stand", level=1)
para(doc,
     "Juriul complet (3–5 membri) ajunge la stand la ora programată. Mentorul/părinții "
     "sunt excluși. Aveți ~15–25 minute: 10 min prezentare + 5–15 min Q&A. Este cea mai "
     "formală variantă de evaluare la stand.",
     size=11, italic=True, color=GREY_LIGHT,
     align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

# ---------- Dispunere fizică ----------
heading(doc, "Dispunere fizică", level=2)
mono_block(doc,
    "                [ DISPLAY OFICIAL 30+60+30 cm ]\n"
    "                            |\n"
    "     [ TV video loop ]      |      [ ROLL-UP 85x200 ]\n"
    "                |           |           |\n"
    "        [ Laptop demo ]     |     [ Flyere / QR ]\n"
    "                 \\          |          /\n"
    "          ALBERT -+---------+---------+- IOAN\n"
    "         (stânga)                       (dreapta)\n"
    "                            |\n"
    "              === JURIUL (3-5 persoane) ===",
    size=9)

para(doc, "", space_after=6)
bullet(doc, "Albert — stânga displayului, zona „viziune/impact\" (context, problemă, Kardashev)")
bullet(doc, "Ioan — dreapta displayului, zona „tehnic/numere\" (lanț energetic, parametri)")
bullet(doc, "Laptop între voi, ușor înclinat spre juriu, demo deja deschis (nu în standby)")
bullet(doc, "TV rulează în fundal cu sunet MUTE în timpul vorbirii — doar imagine ca backdrop")
bullet(doc, "Roll-up lateral vizibil — ca identitate vizuală, nu ca document principal")
bullet(doc, "Flyere pe colțul mesei, la îndemână — oferite la închidere, nu pe parcurs")

# ---------- Flow ----------
heading(doc, "Flow minute-cu-minute", level=2)
flow = [
    ("Minut", "Cine", "Ce face"),
    ("0:00", "Amândoi", "Salut simultan, scurt. „Bună ziua, suntem echipa POLARIS Bears — Albert și Ioan.\""),
    ("0:10", "Albert", "Pornește Varianta A (Inspirațional) sau C (Narativ) — 75–90 sec"),
    ("1:30", "Albert → Ioan", "Tranziție: „Cum ajungem de la Soare la priză? Ioan — tu ai derivat calculele.\""),
    ("1:30", "Ioan", "Plimbă juriul prin display (stânga → dreapta): Dyson Ring → Colectori → L4/L5 → Polare → YBCO. Arată cu mâna pe display, nu citește."),
    ("4:00", "Ioan", "Scoate calculul-cheie: „La 7M km, 622 kW/m². 30 oglinzi × 2,36 × 10⁷ m² × lanț de eficiențe = 1,6 TW.\""),
    ("5:00", "Ioan → Albert", "Tranziție: „Dar ce înseamnă asta pentru lumea reală? Albert.\""),
    ("5:00", "Albert", "Impact + Kardashev + Star Power Grid (surplusul 53%)"),
    ("7:00", "Amândoi", "Demo laptop — animația DiagramZoom. MAX 90 SECUNDE."),
    ("8:30", "Amândoi", "Închide demo-ul. Nu prelungi niciodată."),
    ("8:30", "Amândoi", "Scurt despre echipă: „Am lucrat complementar. Ioan — fizica. Albert — viziune + design. Am verificat fiecare număr împreună.\""),
    ("9:30", "Amândoi", "Închei cu sloganul (vezi mai jos). Tăcere. Pas înapoi."),
]
make_table(doc, flow, [1.5, 3.5, 11.5])
para(doc, "", space_after=8)

# ---------- Dialoguri cheie ----------
heading(doc, "Dialoguri cheie de memorat", level=2)

para(doc, "Tranziția Albert → Ioan:", bold=True, size=10.5, color=NAVY, space_after=2)
quote_block(doc, "„Cum ajungem de la Soare la priză? Ioan — tu ai derivat calculele, te rog.\"", size=10.5)

para(doc, "Tranziția Ioan → Albert:", bold=True, size=10.5, color=NAVY, space_after=2)
quote_block(doc, "„Asta e partea tehnică. Dar ce înseamnă pentru lumea reală? Albert.\"", size=10.5)

para(doc, "Închiderea (amândoi, aliniați, cu sloganul):", bold=True, size=10.5, color=NAVY, space_after=2)
quote_block(doc,
    f"„P.O.L.A.R.I.S. nu e proiectul nostru singur. E contribuția noastră la o conversație "
    f"care durează de 60 de ani — de la Freeman Dyson încoace. {SLOGAN} Mulțumim.\"",
    size=10.5, color=GOLD)

# ---------- Reguli Q&A ----------
heading(doc, "Reguli ferme pentru sesiunea Q&A", level=2)
bullet(doc, "Lăsați juriul să termine întrebarea. Nu întrerupeți niciodată.")
bullet(doc, "Pauză de 1 secundă înainte de răspuns — pare reflexie, nu improvizație.")
bullet(doc, "Cel mai relevant răspunde primul. Fizică → Ioan. Impact → Albert.")
bullet(doc, "Completare, nu contrazicere. „Exact, și aș adăuga…\" — niciodată „nu, de fapt…\"")
bullet(doc, "Nu știi? Spune. „Nu am calculat exact, dar intuiția este X — putem verifica.\"")
bullet(doc, "Folosește displayul ca argument. „Uitați aici, în al treilea pas…\"")
bullet(doc, "Privirea merge către jurorul care a întrebat — nu către celălalt coleg.")

# ---------- Capcane ----------
heading(doc, "Capcane comune de evitat", level=2)
for capcana in [
    "Să începeți cu „Bună, deci proiectul nostru e despre…\" — începeți cu Varianta A sau C, ad litteram.",
    "Să citiți de pe display — displayul e pentru juriu, nu pentru voi.",
    "Să vă uitați unul la celălalt în timpul răspunsurilor — privire directă către jurorul care a întrebat.",
    "Să prelungiți demo-ul peste 90 sec — juriul pierde interesul.",
    "Să începeți cu istoric (Dyson 1960, Kardashev 1964) dacă alegeți Varianta A — acela e element din Varianta C.",
    "Să oferiți flyere în timpul prezentării — distrage jurorul.",
    "Să activați sunetul TV-ului în timpul vorbirii — TV-ul e backdrop vizual.",
]:
    bullet(doc, capcana)

page_break(doc)

# ============== SECȚIUNEA 4 — SCENARIUL B ==============

heading(doc, "4. SCENARIUL B — Evaluare Flash / Juror Solo", level=1)
para(doc,
     "Un singur juror (sau 2 max) se apropie de stand între sloturile programate. Poate a "
     "văzut displayul din mers, poate e curios, poate e în inspecție rapidă. Aveți 3–8 "
     "minute. Nu vă anunță cât stă. Trebuie să comprimați totul fără să păreți grăbiți.",
     size=11, italic=True, color=GREY_LIGHT,
     align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

# ---------- Dispunere ----------
heading(doc, "Dispunere fizică — asimetrică", level=2)
mono_block(doc,
    "         [ DISPLAY / TV cu video loop ]\n"
    "                       |\n"
    "          ALBERT ------+------ IOAN\n"
    "           (unul întâmpină, celălalt pare ocupat)\n"
    "                       |\n"
    "               [ 1 juror solo ]",
    size=9)

para(doc, "", space_after=4)
bullet(doc, "NU amândoi în poziție de atenție simultan — ar fi intimidant pentru un juror solitar")
bullet(doc, "Unul este deja în conversație (sau se preface — face ceva pe laptop)")
bullet(doc, "Celălalt întâmpină cu contact vizual + zâmbet neutru, profesional")
bullet(doc, "Odată ce jurorul se oprește, celălalt se alătură natural")

# ---------- Pitch alternative ----------
heading(doc, "Pitch-ul de deschidere — 2 template-uri alternative (decideți pe loc)", level=2)
para(doc,
     "Ambele template-uri sunt pregătite ad litteram. Alegeți unul pe baza semnalelor "
     "jurorului (vezi „Tipuri de juror\" de mai jos). Template 1 e curat și direct. "
     "Template 2 adaugă un „contract\" de 20 sec care reduce riscul ca jurorul grăbit să "
     "plece înainte ca pitch-ul să înceapă.",
     size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)

# ----- Template 1 -----
heading(doc, "Template 1 — Varianta B comprimată (60 sec, direct)", level=3)
para(doc, "Recomandat pentru: juror vizibil fizician/inginer care vrea cifre rapid.",
     italic=True, size=10, color=GREY_LIGHT, space_after=6)

quote_block(doc,
    "„Doamnelor și domnilor, proiectul P.O.L.A.R.I.S. răspunde la o singură întrebare: "
    "pot legile fizicii cunoscute să susțină un sistem energetic spațial care să livreze "
    "1,6 TW continuu către Pământ?\n\n"
    "Răspunsul nostru: DA.\n\n"
    "30 oglinzi la 7,7 milioane km de Soare, 4 sateliți colectori cu randament 43%, "
    "transfer prin Lagrange L4/L5, 2 stații polare cu cabluri YBCO supraconductoare la "
    "77 K, 800 GW per pol.\n\n"
    "Suntem gata pentru întrebări.\""
)

# ----- Template 2 -----
heading(doc, "Template 2 — Contract 20 sec + pitch 60 sec (80 sec total)", level=3)
para(doc, "Recomandat pentru: juror grăbit sau distras — „contractul\" îl fixează.",
     italic=True, size=10, color=GREY_LIGHT, space_after=6)

para(doc, "Secundele 0–20 — contractul (Albert sau cine e mai aproape):",
     bold=True, size=10.5, color=NAVY, space_after=2)
quote_block(doc,
    "„Bună ziua. Dacă îmi dați 60 de secunde, vă pot arăta cum captăm 48% din "
    "electricitatea globală — de la 7 milioane km distanță de Soare.\"",
    size=10.5)

para(doc, "Secundele 20–80 — pitch-ul (Albert + Ioan alternează):",
     bold=True, size=10.5, color=NAVY, space_after=2)
quote_block(doc,
    "„Umanitatea folosește 3,36 TW continuu, 60% fosili. Panourile clasice — factor de "
    "capacitate 20–50%. Noi livrăm 1,6 TW continuu, 24/7, din spațiu.\n\n"
    "(Ioan, arătând displayul) 30 de oglinzi la 7 milioane km reflectă spre 4 sateliți "
    "colectori → laser verde spre L4 și L5 → laser albastru spre centrale polare la 30 "
    "km altitudine → cablu YBCO de 800 GW la sol. Două poluri = 1,6 TW.\n\n"
    "Fiecare componentă există: Parker Solar Probe, JAXA, James Webb, Project Loon, "
    "CERN. Nu inventăm fizică — integrăm la scară.\""
)

# ---------- Tipuri de juror ----------
heading(doc, "Cum citești jurorul — 4 tipuri, 4 tactici", level=2)
tipuri = [
    ("Tipul jurorului", "Semne", "Tactică + template recomandat"),
    ("Curios", "Se apropie, citește displayul, întreabă „de ce așa?\"",
     "Extindeți pitch-ul. Deschideți laptop dacă cere. Template 2."),
    ("Sceptic", "Stă la distanță, ridică sprâncene la cifre mari",
     "Template 1 + calcule concrete. „Uitați exact cum ajungem la 1,6 TW.\" Nu apărați — demonstrați."),
    ("Grăbit", "Privește ceasul, stă 2 min max",
     "Template 1 comprimat la 40 sec. Închideți: „Puteți reveni la ora X pentru demo complet?\""),
    ("Tăcut", "Tăcere prelungită după pitch",
     "Nu umpleți tăcerea aleator. Întrebați: „Ce parte vă interesează cel mai mult?\""),
]
make_table(doc, tipuri, [3.0, 5.5, 8.0])

# ---------- Ce să NU faceți ----------
heading(doc, "Ce să NU faceți în Scenariul B", level=2)
for reg in [
    "Pitch-ul formal de 10 min — jurorul nu a venit pentru asta.",
    "Ambii să vorbiți simultan sau să vă completați fiecare propoziție — pare repetat/script.",
    "Demo pe laptop dacă jurorul e clar grăbit — doar la semnal explicit („pot să văd?\").",
    "Să cereți feedback la final — „cum vi s-a părut?\" nu e rolul jurorului informal.",
    "Să oferiți mingi antistress jurorului — pare comercial; rezervate vizitatorilor non-juriu.",
]:
    bullet(doc, reg)

# ---------- Cârlige intelectuale ----------
heading(doc, "Cârlige intelectuale de închidere (conversie flash → formal)", level=2)
para(doc,
     "Obiectivul Scenariului B nu e să luați punctaj pe loc. Este să faceți jurorul să "
     "vrea să „revină cu colegii\" sau să „audă mai mult la prezentarea formală\". "
     "Pentru asta, lăsați un cârlig la final:",
     size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

para(doc, "Cârlig Tipul 1 — Paradox:", bold=True, size=10.5, color=NAVY, space_after=2)
quote_block(doc,
    "„Curios este că laserul de la 150 milioane km ajunge cu precizie de 1 km — exact "
    "cât are iris-ul nostru captor. Vă putem explica cum, dacă vreți.\"", size=10.5)

para(doc, "Cârlig Tipul 2 — Cifră neașteptată:", bold=True, size=10.5, color=NAVY, space_after=2)
quote_block(doc,
    "„Apropo, break-even-ul de CO₂ al proiectului e de doar 5 ani. Dacă doriți, am "
    "calculul complet.\"", size=10.5)

para(doc, "Cârlig Tipul 3 — Referire la juror:", bold=True, size=10.5, color=NAVY, space_after=2)
quote_block(doc,
    "„Dacă lucrați cu fizică aplicată, probabil vă veți întreba de ce am ales YBCO și "
    "nu NbTi. Răspunsul e interesant.\"", size=10.5)

# ---------- Tranziția B → A ----------
heading(doc, "Tranziția B → A (dacă flash devine formal)", level=2)
para(doc,
     "Dacă jurorul solo stă peste 5 minute și începe întrebări tehnice profunde, "
     "sunteți în Scenariul A improvizat. Semnale:",
     size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
bullet(doc, "Scoate pix/tabletă să-și ia notițe")
bullet(doc, "Ajung alți membri ai juriului")
bullet(doc, "Începe să comenteze pentru sine („interesant… interesant…\")")
para(doc,
     "Când vedeți semnalele, treceți la modul A: poziționare simetrică, dispunere "
     "clasică, structură de 10 min (Varianta A sau C ca deschidere nouă).",
     size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)

page_break(doc)

# ============== SECȚIUNEA 5 — INTEGRAREA MATERIALELOR ==============

heading(doc, "5. Integrarea materialelor de stand", level=1)
integrare = [
    ("Material", "Scenariul A (formal)", "Scenariul B (flash)"),
    ("TV cu video loop",
     "Rulează în fundal, MUTE în timpul vorbirii; backdrop vizual de autoritate",
     "Cârlig vizual principal — ce atrage jurorul la stand; volum discret"),
    ("Display oficial ONCS\n(30+60+30 cm)",
     "Suport principal; Ioan plimbă juriul pe el (min 1:30–5:00)",
     "Referință rapidă — arăt cu mâna pe el în pitch"),
    ("Roll-up marketing\n(85×200 cm)",
     "Lateral, identitate vizuală — NU înlocuiește displayul",
     "Identitate vizuală brand; QR code vizibil"),
    ("Laptop + demo",
     "Obligatoriu la minutul 7:00, MAX 90 sec",
     "Doar la cerere explicită a jurorului"),
    ("Flyere",
     "Pe colțul mesei; oferite la închidere (nu pe parcurs)",
     "Oferite la final cu QR — take-away pentru revizitare"),
    ("Mingi antistress",
     "NU în timpul evaluării",
     "NU jurorului; pentru vizitatori non-juriu"),
    ("Website + QR",
     "Menționat la închidere",
     "Focus principal în flyer-ul take-away"),
]
make_table(doc, integrare, [3.5, 6.0, 6.5])

page_break(doc)

# ============== SECȚIUNEA 6 — MATRICEA DE PREGĂTIRE ==============

heading(doc, "6. Matricea de pregătire", level=1)
matrice = [
    ("Element", "Scenariul A (Formal)", "Scenariul B (Flash)"),
    ("Pitch de deschidere", "Varianta A (90 sec) sau C (75 sec) — ad litteram",
     "Template 1 (60 sec) sau Template 2 (80 sec) — decideți pe loc"),
    ("Prezentare principală", "10 min structurată pe flow minute-cu-minute", "90 sec comprimată"),
    ("Demo laptop", "Obligatoriu, 90 sec max", "Doar la cerere explicită"),
    ("Q&A", "5–15 min; acoperire pe toate cele 54 întrebări pregătite", "2–5 min focusat pe curiozitate"),
    ("Poziționare fizică", "Simetrică, formală, ambii în atenție", "Asimetrică, relaxată, unul întâmpină"),
    ("Închidere", "„Mulțumim\" + slogan", "Cârlig intelectual + invitație revenire"),
    ("Repetiții necesare", "Minim 5× cronometrat (cu toate tranzițiile)",
     "Minim 10× — automatism total pe ambele template-uri"),
]
make_table(doc, matrice, [3.8, 6.0, 6.2])
para(doc, "", space_after=12)

# ============== SECȚIUNEA 7 — REGULA UNICĂ + CHECKLIST ==============

heading(doc, "7. Regula unică pentru amândouă scenariile", level=1)
quote_block(doc,
    "Juriul nu notează proiectul. Notează impresia pe care voi, ca echipă, ați reușit "
    "s-o creați despre proiect. Un proiect bun prezentat slab pierde. Un proiect bun "
    "prezentat excelent câștigă inclusiv puncte pe care nu le merită strict tehnic. "
    "La stand, VOI sunteți interfața. Calculele sunt pe display. Numerele sunt în "
    "fișă. Dar încrederea se transmite prin voce, privire, calm și claritate.",
    size=11.5, color=NAVY)

heading(doc, "Checklist final înainte de ziua competiției", level=2)
for item in [
    "Varianta A (Inspirațional) — exersată cronometrat de cel puțin 5 ori.",
    "Varianta C (Narativ) — exersată cronometrat de cel puțin 3 ori.",
    "Template 1 + Template 2 (Scenariul B) — automatism total, minim 10× fiecare.",
    "Cele 54 întrebări Q&A parcurse individual de fiecare membru.",
    "Demo laptop DiagramZoom funcționează — backup video pe telefon + stick USB.",
    "TV conectat la laptop (cablu HDMI + adaptor) — testat la fața locului.",
    "Flyere în stoc vizibil dar nu intruziv; mingi antistress în cutie, nu pe masă.",
    "Toate numerele cheie memorate: 1,6 TW · 48% · 7M km · 30 oglinzi · 93 K · 800 GW.",
    "Cârlige intelectuale pregătite (Paradox / Cifră / Referire la juror).",
    "Pregătit răspunsul de rezervă: „Nu știu. Am o intuiție, dar nu am calcul exact.\"",
    "Poziționarea fizică discutată: Albert-stânga, Ioan-dreapta — nu inversați.",
]:
    bullet(doc, item)

page_break(doc)

# ============== PAGINA FINALĂ — SEMNĂTURA ==============

para(doc, "", space_after=80)
para(doc, "„", bold=True, size=48, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_after=0)
para(doc, SLOGAN, bold=True, italic=True, size=26, color=GOLD_BRIGHT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "„", bold=True, size=48, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_after=40)

para(doc, "Echipa POLARIS Bears",
     bold=True, size=14, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "Albert David OLARIU  ·  Ioan Cristian CHELARU",
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "Mentor: Prof. Dr. Francisc Dionisie Aaron",
     italic=True, size=10, color=GREY_LIGHT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "Laureat al Premiului Academiei Române (2002) — „Mecanica Analitică\"",
     italic=True, size=9, color=GREY_VLIGHT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

para(doc, "ONCS 2026  ·  Secțiunea C — Tehnologia Informației  ·  Juniori",
     size=10, color=GREY_LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "polaris-bears.ro",
     bold=True, size=10, color=CYAN,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

# ============== SAVE ==============

doc.save(OUTPUT)
print(f"OK: salvat la {OUTPUT}")
print(f"Format: A4 portrait (210 x 297 mm)")
print(f"Sectiuni: 7 + cover + semnatura finala")
print(f"Pitch-uri: 3 variante (A, B, C) + 2 template-uri Scenariul B")
print(f"Slogan: „{SLOGAN}\" — pe cover + semnatura finala")

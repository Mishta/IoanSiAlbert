#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_oncs_doc.py — Document ONCS 2026 pentru proiectul P.O.L.A.R.I.S.

Generează: docs/POLARIS_ONCS2026_Inscriere.docx
Run: py -3.14 generate_oncs_doc.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Helpers ───────────────────────────────────────────────────────────────────

def add_h(doc, text, level=1, before=10, after=4):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(before)
    h.paragraph_format.space_after = Pt(after)
    return h

def add_p(doc, text, italic=False, bold=False, size=11, after=6,
          indent=0, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, items, indent=0.35, size=11, after=3):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.left_indent = Inches(indent)
        run = p.add_run(f"\u2022 {item}")
        run.font.size = Pt(size)

def add_numbered(doc, items, indent=0.35, size=11, after=3):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.left_indent = Inches(indent)
        run = p.add_run(f"{i}. {item}")
        run.font.size = Pt(size)

def set_col_width(table, col_widths_inches):
    """Setează lățimile coloanelor unui tabel (în inches)."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths_inches):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(col_widths_inches[i] * 1440)))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

def table_row(row, values, bold_first=False, size=10):
    for i, (cell, val) in enumerate(zip(row.cells, values)):
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(size)
                r.bold = (bold_first and i == 0)


# ── Build ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Margini pagină
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.18)
        sec.right_margin  = Inches(1.18)

    # Font implicit
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════════════
    # COPERTĂ
    # ═══════════════════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("P.O.L.A.R.I.S.")
    r.bold = True; r.font.size = Pt(32)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_sub.add_run("Polar Orbital Laser Architecture for Renewable Interconnected Solar-power")
    r.italic = True; r.font.size = Pt(12)

    doc.add_paragraph()

    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep.add_run("─" * 60).font.size = Pt(9)

    doc.add_paragraph()

    p_ev = doc.add_paragraph()
    p_ev.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_ev.add_run("Olimpiada Națională de Cercetare Științifică — ONCS 2026")
    r.bold = True; r.font.size = Pt(13)

    p_sec = doc.add_paragraph()
    p_sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sec.add_run("Secțiunea A — Științe Exacte").font.size = Pt(11)

    for _ in range(2):
        doc.add_paragraph()

    p_team = doc.add_paragraph()
    p_team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_team.add_run("Ioan Cristian CHELARU  ·  Albert David OLARIU")
    r.bold = True; r.font.size = Pt(12)

    p_ment = doc.add_paragraph()
    p_ment.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_ment.add_run(
        "Coordonator: Prof. Aaron FRANCISC\n"
        "ICHB — Liceul Teoretic Internațional de Informatică București"
    )
    r.font.size = Pt(10)

    for _ in range(3):
        doc.add_paragraph()

    p_year = doc.add_paragraph()
    p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_year.add_run("Aprilie 2026")
    r.font.size = Pt(11); r.bold = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CUPRINS
    # ═══════════════════════════════════════════════════════════════════════════
    h_toc = doc.add_heading("CUPRINS", level=1)
    h_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_toc.paragraph_format.space_after = Pt(12)

    toc = [
        ("1", "Date de identificare", False),
        ("2", "Rezumat proiect (max. 250 cuvinte)", False),
        ("3", "Descrierea detaliată a proiectului", False),
        ("3.1", "Scop", True),
        ("3.2", "Obiective", True),
        ("3.3", "Problema identificată și stadiul actual în domeniu", True),
        ("3.4", "Etape parcurse", True),
        ("3.5", "Metode folosite", True),
        ("3.6", "Date experimentale și calcule", True),
        ("3.7", "Concluzii", True),
        ("3.8", "Bibliografie", True),
        ("3.9", "Anexe", True),
    ]
    for num, title, sub in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3 if sub else 0)
        r = p.add_run(f"{num}.{'  ' if not sub else '  '}{title}")
        r.font.size = Pt(10 if sub else 11)
        r.bold = not sub

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 1. DATE DE IDENTIFICARE
    # ═══════════════════════════════════════════════════════════════════════════
    add_h(doc, "1. Date de identificare", level=1, before=0, after=8)

    id_table = doc.add_table(rows=7, cols=2)
    id_table.style = 'Table Grid'

    id_data = [
        ("Acronimul proiectului",
         "P.O.L.A.R.I.S.\n"
         "(Polar Orbital Laser Architecture for Renewable Interconnected Solar-power)"),
        ("Categorie",
         "[De completat conform categoriei ONCS 2026]"),
        ("Secțiune",
         "Secțiunea A — Științe Exacte"),
        ("Mentorul echipei",
         "Prof. Aaron FRANCISC\n"
         "ICHB — Liceul Teoretic Internațional de Informatică București\n"
         "Tel.: [De completat]"),
        ("Elev 1 — Ioan Cristian CHELARU",
         "Clasa: [De completat] · Unitatea școlară: ICHB\n"
         "Rol: Modelare vizuală, scala Kardashev, rețeaua Star Power Grid, prezentare multimedia"),
        ("Elev 2 — Albert David OLARIU",
         "Clasa: [De completat] · Unitatea școlară: ICHB\n"
         "Rol: Arhitectura sistem, calcule energetice, supraconductori YBCO, dinamică orbitală"),
        ("Colaboratori",
         "Daniel-Justinian ZELENSCHI — Ing. chimic (UCL + Melbourne), PhD(c) AI, "
         "NASA Ames Research Center; Space Settlement Contest — Premiul II\n"
         "Carlo Emilio MONTANARI — Fizician aplicat, Doctorand Dinamica Fasciculelor, "
         "CERN / Univ. Bologna; consultanță: dinamica fasciculelor laser de înaltă precizie"),
    ]

    for i, (k, v) in enumerate(id_data):
        row = id_table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True; r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

    set_col_width(id_table, [1.8, 5.2])
    doc.add_paragraph()

    add_p(doc,
          "Notă: Câmpurile marcate [De completat] trebuie completate de mentor înainte de depunere.",
          italic=True, size=9, color=(0xaa, 0x44, 0x00), after=4)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 2. REZUMAT PROIECT
    # ═══════════════════════════════════════════════════════════════════════════
    add_h(doc, "2. Rezumat proiect", level=1, before=0, after=6)
    add_p(doc, "Cerință Regulament: maximum 250 de cuvinte.",
          italic=True, size=9, color=(0x55, 0x55, 0x99), after=10)

    # ── 2A. Varianta echipei ──────────────────────────────────────────────────
    add_h(doc, "Varianta A — redactată de echipă (de selectat sau modificat)", level=2,
          before=4, after=4)

    varianta_a = (
        "Proiectul propune un model teoretic al unui Sistem Energetic Spațial (SES) capabil să capteze "
        "energia solară în proximitatea Soarelui și să o transmită pe Pământ. Sistemul utilizează un Dyson "
        "swarm de 30 de oglinzi orbitale de 1 km diametru fiecare, plasate la aproximativ 7 milioane km de "
        "Soare, care reflectă radiația solară către patru sateliți colectori. Aceștia transformă lumina în "
        "fascicule laser cu randament de 43%, direcționate către noduri energetice amplasate în punctele "
        "Lagrange L4 și L5. Centralele orbitale polare, susținute la altitudini de ~30 km prin baloane cu "
        "heliu, convertesc energia laser în electricitate și o transmit la sol prin cabluri supraconductoare "
        "YBCO, fiecare capabil de 800 GW, rezultând o putere totală de ~1,6 TW. Proiectul examinează "
        "fezabilitatea fizică a acestei infrastructuri, incluzând calculul fluxului energetic, limitările "
        "impuse de densitatea de curent a cablurilor YBCO și flotabilitatea baloanelor. De asemenea, propune "
        "extensia vizionar\u0103 \u201eStar Power Grid\u201d, o re\u021bea orbital\u0103 de microunde pentru alimentarea regiunilor "
        "izolate. Analiza arată că, chiar cu pierderile de conversie și limitările tehnologice actuale, "
        "sistemul poate furniza aproape jumătate din necesarul global de energie electrică, respectând legile "
        "fundamentale ale fizicii. Proiectul combină fizica clasică, supraconductivitatea, criogenia, "
        "ingineria orbitală și power beaming, oferind o viziune interdisciplinară pentru viitorul energiei "
        "spațiale, cu relevanță educațională și științifică."
    )
    wc_a = len(varianta_a.split())

    p_a = doc.add_paragraph(varianta_a)
    p_a.paragraph_format.space_after = Pt(4)
    p_a.paragraph_format.first_line_indent = Inches(0.3)
    p_a.runs[0].font.size = Pt(11)

    p_wca = doc.add_paragraph()
    p_wca.paragraph_format.space_after = Pt(12)
    r_wca = p_wca.add_run(f"[{wc_a} cuvinte — limita: 250]")
    r_wca.italic = True; r_wca.font.size = Pt(9)
    r_wca.font.color.rgb = RGBColor(*(0x33, 0x88, 0x33) if wc_a <= 250 else (0xcc, 0x22, 0x22))

    # ── 2B. Varianta Claude ───────────────────────────────────────────────────
    add_h(doc, "Varianta B — propusă de asistentul de cercetare", level=2,
          before=4, after=4)

    varianta_b_blocks = [
        "P.O.L.A.R.I.S. propune un model teoretic pentru un Sistem Energetic Spațial (SES) capabil să "
        "capteze energia solară în vecinătatea Soarelui și să o livreze la rețelele terestre, eliminând "
        "cele două limitări fundamentale ale energiei solare clasice: atenuarea atmosferică și ciclul "
        "zi–noapte.",

        "Arhitectura sistemului se articulează în patru componente interconectate: (1) un Dyson swarm de "
        "30 de oglinzi orbitale cu diametrul de 1 km, amplasate la ~7 milioane km de Soare; (2) patru "
        "sateliți colectori ce convertesc radiația în fascicule laser cu randament de 43%; (3) noduri "
        "energetice în punctele Lagrange L4 și L5 ale sistemului Soare–Pământ; (4) două centrale orbitale "
        "polare susținute la ~30 km altitudine prin baloane cu heliu, conectate la sol prin cabluri "
        "supraconductoare YBCO de 800 GW fiecare. Puterea totală livrată este de ~1,6 TW — echivalentul "
        "a circa 45% din consumul electric global actual.",

        "Proiectul demonstrează coerența fizică a lanțului energetic, cuantifică pierderile la fiecare "
        "etapă de conversie și analizează limitele tehnologice actuale: densitatea de curent YBCO, "
        "cerințele criogenice și flotabilitatea baloanelor (~11 tone per centrală). Este introdusă și "
        "extensia \u201eStar Power Grid\u201d \u2014 o re\u021bea orbital\u0103 de microunde care ancoreaz\u0103 P.O.L.A.R.I.S. \u00een "
        "contextul tranziției Kardashev de la civilizația de tip 0 la tip I.",

        "Lucrarea combină fizica clasică, supraconductivitatea, ingineria orbitală și power beaming "
        "\u00eentr-un cadru interdisciplinar care distinge riguros \u00eentre \u201eposibil fizic\u201d \u0219i \u201eposibil tehnologic\u201d.",
    ]
    varianta_b_full = " ".join(varianta_b_blocks)
    wc_b = len(varianta_b_full.split())

    for block in varianta_b_blocks:
        p_b = doc.add_paragraph(block)
        p_b.paragraph_format.space_after = Pt(5)
        p_b.paragraph_format.first_line_indent = Inches(0.3)
        p_b.runs[0].font.size = Pt(11)

    p_wcb = doc.add_paragraph()
    p_wcb.paragraph_format.space_after = Pt(12)
    r_wcb = p_wcb.add_run(f"[{wc_b} cuvinte — limita: 250]")
    r_wcb.italic = True; r_wcb.font.size = Pt(9)
    r_wcb.font.color.rgb = RGBColor(*(0x33, 0x88, 0x33) if wc_b <= 250 else (0xcc, 0x22, 0x22))

    # ── Tabel comparativ ──────────────────────────────────────────────────────
    add_h(doc, "Comparație variante — criterii de selecție", level=2, before=4, after=6)

    comp_t = doc.add_table(rows=8, cols=3)
    comp_t.style = 'Table Grid'
    comp_rows = [
        ("Criteriu", "Varianta A (echipă)", "Varianta B (asistent)"),
        ("Ton", "Tehnic, dens, narativ", "Structurat, concis, cu enumerare"),
        ("Structură", "Paragraf unic curgător", "4 paragrafe logice distincte"),
        ("Cifre cheie", "43%, 800 GW, 1,6 TW", "43%, 800 GW, 1,6 TW, 45% global"),
        ("Star Power Grid", "Menționat, explicat", "Legat explicit de Scala Kardashev"),
        ("Interdisciplinaritate", "Menționată la final", "Integrată pe tot parcursul"),
        ("Nr. cuvinte", f"{wc_a} cuvinte", f"{wc_b} cuvinte"),
        ("Recomandare", "Stil narativ, mai accesibil", "Structură clară, impact tehnic mai mare"),
    ]
    for i, row_data in enumerate(comp_rows):
        row = comp_t.rows[i]
        for j, (cell, val) in enumerate(zip(row.cells, row_data)):
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.bold = (i == 0)
    set_col_width(comp_t, [1.8, 2.7, 2.7])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 3. DESCRIEREA DETALIATĂ
    # ═══════════════════════════════════════════════════════════════════════════
    add_h(doc, "3. Descrierea detaliată a proiectului", level=1, before=0, after=4)
    add_p(doc, "Maxim 10 pagini — doar activitățile desfășurate în proiect.",
          italic=True, size=9, color=(0x88, 0x88, 0x88), after=10)

    # 3.1
    add_h(doc, "3.1 Scop", level=2)
    add_p(doc,
          "Scopul proiectului P.O.L.A.R.I.S. este modelarea teoretică și demonstrarea coerenței fizice "
          "a unui Sistem Energetic Spațial (SES) capabil să capteze energie solară în proximitatea "
          "Soarelui și să o transmită la nivelul rețelei terestre, cu o putere totală de ~1,6 TW. "
          "Proiectul vizează identificarea limitelor tehnologice actuale și evaluarea fezabilității "
          "pe termen lung a unei astfel de infrastructuri energetice globale.")

    # 3.2
    add_h(doc, "3.2 Obiective", level=2)
    add_numbered(doc, [
        "Modelarea fluxului energetic complet: de la colectarea circumsolară la livrarea la sol.",
        "Dimensionarea componentelor cheie: oglinzi Dyson swarm, sateliți colectori, cabluri YBCO, baloane.",
        "Evaluarea randamentelor de conversie la fiecare etapă a lanțului energetic.",
        "Analiza limitelor tehnologice actuale: supraconductori YBCO, criogenie, flotabilitate.",
        "Demonstrarea coerenței fizice a sistemului raportat la legile fundamentale ale fizicii.",
        "Introducerea extensiei Star Power Grid și poziționarea în contextul Scării Kardashev.",
    ])
    doc.add_paragraph()

    # 3.3
    add_h(doc, "3.3 Problema identificată și stadiul actual în domeniu", level=2)
    add_p(doc,
          "Criza energetică globală este definită de trei provocări structurale: cererea în creștere "
          "(~3,36 TW echivalent electric în 2024, conform IEA), intermitența surselor regenerabile "
          "terestre și limitările impuse de atmosferă și ciclul zi–noapte. Sursele solare terestre "
          "pierd 20–30% din radiație prin atenuare atmosferică și sunt inactive ~50% din timp.")
    add_p(doc,
          "Conceptul Space-Based Solar Power (SBSP) a fost introdus de Glaser (1968) și dezvoltat "
          "prin programele NASA NIAC (SPS-ALPHA, Mankins 2011) și JAXA SSPS. Demonstrațiile la "
          "scară mică au validat transmisia wireless prin microunde (Brown, 1984; JAXA). P.O.L.A.R.I.S. "
          "extinde paradigma SBSP prin plasarea colectorilor în proximitatea Soarelui (~7M km, validat "
          "de traiectoria sondei Parker Solar Probe, NASA 2018), maximizând irradianța disponibilă.")

    # 3.4
    add_h(doc, "3.4 Etape parcurse", level=2)
    stages = [
        ("Etapa 1 — Documentare",
         "Studiul literaturii: Glaser (1968), NASA NIAC, JAXA SSPS, Parker Solar Probe, proprietăți YBCO."),
        ("Etapa 2 — Arhitectură sistem",
         "Definirea componentelor: Dyson swarm (30 oglinzi × 1 km diametru), 4 sateliți colectori, "
         "2 noduri Lagrange (L4/L5), 2 centrale polare (~30 km altitudine)."),
        ("Etapa 3 — Calcule energetice",
         "Modelarea fluxului energetic complet, calculul randamentelor cumulate, "
         "dimensionarea cablurilor YBCO și flotabilitatea baloanelor."),
        ("Etapa 4 — Analiză limite tehnologice",
         "Evaluarea masei totale per centrală (~11 tone), limitei de curent YBCO (~500 A/mm²) "
         "și cerințelor criogenice la altitudine."),
        ("Etapa 5 — Star Power Grid",
         "Conceptualizarea extensiei cu transmisie orbitală prin microunde și poziționarea "
         "pe Scala Kardashev (tranziție tip 0 → tip I)."),
        ("Etapa 6 — Prezentare multimedia",
         "Elaborarea prezentării P.O.L.A.R.I.S. (19 slide-uri) cu vizualizări generate AI "
         "(Kling AI, Google Gemini) și design Celestial Command."),
    ]
    for stage, desc in stages:
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_after = Pt(4)
        p_s.paragraph_format.left_indent = Inches(0.3)
        r_s = p_s.add_run(f"{stage}: ")
        r_s.bold = True; r_s.font.size = Pt(11)
        p_s.add_run(desc).font.size = Pt(11)
    doc.add_paragraph()

    # 3.5
    add_h(doc, "3.5 Metode folosite", level=2)
    add_p(doc,
          "Proiectul a utilizat metode teoretice și de modelare numerică, fără experimente fizice "
          "directe. Metodele principale sunt:")
    add_bullet(doc, [
        "Calcul analitic al fluxului energetic (irradianță solară, randamente de reflexie, "
        "conversie laser, transmisie supraconductoare).",
        "Modelare geometrică a orbitelor și unghiurilor de reflexie pentru Dyson swarm.",
        "Analiză structurală a flotabilității baloanelor (principiul lui Arhimede, densitate "
        "He vs. aer la ~30 km altitudine).",
        "Consultanță externă: Carlo Emilio MONTANARI (CERN / Univ. Bologna) — dinamica "
        "fasciculelor laser de înaltă precizie (transfer de la LHC la P.O.L.A.R.I.S.).",
        "Documentare și sinteză bibliografică (12 surse academice, stil Harvard).",
        "Vizualizare și prezentare: Kling AI (generare imagini), Google Gemini, "
        "Microsoft PowerPoint 2019 / 365.",
    ])
    doc.add_paragraph()

    # 3.6
    add_h(doc, "3.6 Date experimentale și calcule", level=2)
    add_p(doc,
          "Calculele au fost realizate pe baza datelor fizice măsurate și publicate. "
          "Parametrii principali și rezultatele sunt sintetizate mai jos:")

    calc_t = doc.add_table(rows=10, cols=3)
    calc_t.style = 'Table Grid'
    calc_rows = [
        ("Parametru", "Valoare", "Sursă"),
        ("Irradianță solară la 1 UA", "1.361 W/m²", "NASA GSFC"),
        ("Irradianță la ~0,047 UA (7M km)", "~612.500 W/m²", "Calcul (legea 1/r²)"),
        ("Suprafață totală oglinzi Dyson", "30 × π × (500 m)² ≈ 23,6 km²", "Proiect"),
        ("Putere captată (η_ref = 85%)", "~12.270 GW", "Calcul"),
        ("Randament conversie laser", "43%", "Proiect / lit."),
        ("Randament transmisie Lagrange→Pol", "~90%", "Estimare"),
        ("Putere per cablu YBCO", "800 GW", "Proiect"),
        ("Putere totală livrată la sol", "~1,6 TW", "Calcul"),
        ("Masă totală per centrală polară", "~11 tone", "Estimare proiect"),
    ]
    for i, rd in enumerate(calc_rows):
        row = calc_t.rows[i]
        for j, (cell, val) in enumerate(zip(row.cells, rd)):
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.bold = (i == 0)
    set_col_width(calc_t, [2.5, 2.1, 2.1])
    doc.add_paragraph()

    # 3.7
    add_h(doc, "3.7 Concluzii", level=2)
    add_bullet(doc, [
        "Sistemul P.O.L.A.R.I.S. este fizic fezabil — respectă legile fundamentale ale fizicii "
        "și ecuațiile de transfer energetic.",
        "Puterea totală de ~1,6 TW corespunde cu ~45% din consumul electric global actual "
        "(IEA 2024 — 3,36 TW).",
        "Limitele identificate sunt tehnologice, nu fizice: proprietățile actuale ale YBCO, "
        "criogenia la altitudine și materialele pentru oglinzi.",
        "Star Power Grid reprezintă extensia naturală a sistemului și marchează tranziția "
        "spre civilizația de tip I pe Scala Kardashev.",
        "Proiectul demonstrează că soluțiile energetice ale viitorului necesită o abordare "
        "interdisciplinară, integrând fizică, inginerie și viziune strategică.",
    ])
    doc.add_paragraph()

    # 3.8 Bibliografie
    doc.add_page_break()
    add_h(doc, "3.8 Bibliografie", level=2, before=0)
    add_p(doc, "Stilul de citare: Harvard (autor, an, titlu, sursă, DOI/URL).",
          italic=True, size=9, color=(0x77, 0x77, 0x77), after=8)

    references = [
        ("01", "Glaser, P.E., 1968. Power from the Sun: Its Future. Science, 162(3856), pp.857–861.",
         "https://doi.org/10.1126/science.162.3856.857"),
        ("02", "International Energy Agency (IEA), 2024. Electricity 2024: Analysis and forecast to 2026.",
         "https://www.iea.org/reports/electricity-2024"),
        ("03", "Japan Aerospace Exploration Agency (JAXA), n.d. Space Solar Power Systems (SSPS).",
         "https://www.kenkai.jaxa.jp/eng/research/ssps/"),
        ("04", "Mankins, J.C., 2011. SPS-ALPHA: The First Practical Solar Power Satellite via "
               "Arbitrarily Large Phased Array. NASA NIAC Phase I Report.",
         "https://www.nasa.gov/sites/default/files/files/niac_2011_phaseI_mankins.pdf"),
        ("05", "NASA, n.d. Space-Based Solar Power. NASA Innovative Advanced Concepts (NIAC).",
         "https://www.nasa.gov/directorates/space-tech/niac/space-based-solar-power/"),
        ("06", "NASA Goddard Space Flight Center, n.d. Total Solar Irradiance.",
         "https://solarscience.msfc.nasa.gov/"),
        ("07", "NASA, n.d. What are Lagrange Points? NASA Science — Solar System Exploration.",
         "https://solarsystem.nasa.gov/resources/754/what-are-lagrange-points/"),
        ("08", "NASA, 2018. Parker Solar Probe Mission Overview.",
         "https://www.nasa.gov/content/goddard/parker-solar-probe"),
        ("09", "Brown, W.C., 1984. The History of Power Transmission by Radio Waves. "
               "IEEE Transactions on Microwave Theory and Techniques, 32(9), pp.1230–1242.", ""),
        ("10", "Shinohara, N., 2014. Wireless Power Transfer via Radiowaves. London: ISTE Ltd.", ""),
        ("11", "CERN, 2021. High-temperature superconductors for future accelerators.",
         "https://home.cern/science/engineering/superconductivity"),
        ("12", "Larbalestier, D., Gurevich, A., Feldmann, D.M. și Polyanskii, A., 2001. "
               "High-Tc superconducting materials for electric power applications. "
               "Nature, 414(6861), pp.368–377.",
         "https://doi.org/10.1038/35104654"),
    ]

    for num, text, url in references:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_after = Pt(5)
        p_ref.paragraph_format.left_indent = Inches(0.4)
        p_ref.paragraph_format.first_line_indent = Inches(-0.4)
        r_num = p_ref.add_run(f"[{num}] ")
        r_num.bold = True; r_num.font.size = Pt(10)
        r_txt = p_ref.add_run(text)
        r_txt.font.size = Pt(10)
        if url:
            r_url = p_ref.add_run(f"\n        {url}")
            r_url.font.size = Pt(9)
            r_url.font.color.rgb = RGBColor(0x00, 0x44, 0x99)

    doc.add_paragraph()

    # 3.9 Anexe
    add_h(doc, "3.9 Anexe", level=2, before=8)
    add_bullet(doc, [
        "Anexa 1: Prezentarea multimedia P.O.L.A.R.I.S. (19 slide-uri) — fișier PPTX atașat.",
        "Anexa 2: Schema arhitecturală completă a sistemului SES.",
        "Anexa 3: Diagrama fluxului energetic cu randamente detaliate pe etape.",
        "Anexa 4: Grafic comparativ Scala Kardashev — poziționarea P.O.L.A.R.I.S.",
    ])

    # ═══════════════════════════════════════════════════════════════════════════
    # SALVARE
    # ═══════════════════════════════════════════════════════════════════════════
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "output", "POLARIS_ONCS2026_Inscriere.docx")
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    print(f"Document generat: {out}")
    return out


if __name__ == "__main__":
    build()
